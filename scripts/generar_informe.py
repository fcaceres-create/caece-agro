"""
Genera el informe TFI académico final del proyecto AgroSmart en docs/informe.docx.

El script es reproducible: cualquier modificación del contenido (texto, métricas,
figuras) se actualiza acá y se regenera el .docx con `python scripts/generar_informe.py`.
La fuente del contenido es la bitácora docs/recuperacion_de_datos.md (14 secciones)
y los artefactos del repo (data/modelos/reporte_entrenamiento.json, figuras del EDA).

NOTA PARA LA ENTREGA FINAL:
  El capítulo 7 incluye un placeholder para una captura de pantalla de la app
  Streamlit. Antes de entregar el informe, abrir el .docx en Word y reemplazar
  el placeholder con la imagen real (instrucción explícita en el cuerpo del doc).

Para correr:
    python scripts/generar_informe.py
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------
RAIZ = Path(__file__).resolve().parents[1]
DOCS_DIR = RAIZ / "docs"
FIGURAS_DIR = DOCS_DIR / "figuras"
INFORME_PATH = DOCS_DIR / "informe.docx"
METRICAS_PATH = RAIZ / "data" / "modelos" / "reporte_entrenamiento.json"

# ---------------------------------------------------------------------
# Constantes de estilo
# ---------------------------------------------------------------------
FUENTE_CUERPO = "Calibri"
FUENTE_CODIGO = "Consolas"
COLOR_TITULO = RGBColor(0x1F, 0x4E, 0x79)  # azul oscuro académico
COLOR_SUBTITULO = RGBColor(0x2E, 0x75, 0xB6)
COLOR_CAPTION = RGBColor(0x59, 0x59, 0x59)
ANCHO_FIGURA = Inches(6.0)


# =====================================================================
# Helpers de bajo nivel (XML directo para cosas que python-docx no expone)
# =====================================================================
def _set_cell_border(cell, **kwargs) -> None:
    """Aplica bordes simples a una celda. kwargs: top, bottom, left, right."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "808080")
            borders.append(border)
    tc_pr.append(borders)


def _shade_cell(cell, color_hex: str) -> None:
    """Sombrea una celda con color hex (ej 'D9E2F3')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _agregar_campo(parrafo, instruccion: str, texto_placeholder: str) -> None:
    """Inserta un campo OOXML (ej. TOC, PAGE, NUMPAGES).

    Word evalúa estos campos al abrir el documento. python-docx no
    sabe renderizarlos, por eso se inyecta el XML directo.
    """
    run = parrafo.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruccion
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = texto_placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(el)


def _saltar_pagina(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# =====================================================================
# Configuración global del documento
# =====================================================================
def configurar_estilos(doc: Document) -> None:
    """Define fuentes, tamaños y interlineado para los estilos base."""
    estilos = doc.styles

    normal = estilos["Normal"]
    normal.font.name = FUENTE_CUERPO
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for nivel, color, tam in (
        ("Heading 1", COLOR_TITULO, 18),
        ("Heading 2", COLOR_SUBTITULO, 14),
        ("Heading 3", COLOR_SUBTITULO, 12),
    ):
        h = estilos[nivel]
        h.font.name = FUENTE_CUERPO
        h.font.size = Pt(tam)
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True

    # Estilo para bloques de código (definido a mano para no chocar con built-ins).
    codigo = estilos.add_style("CodigoMono", WD_STYLE_TYPE.PARAGRAPH)
    codigo.font.name = FUENTE_CODIGO
    codigo.font.size = Pt(9)
    codigo.paragraph_format.left_indent = Cm(0.5)
    codigo.paragraph_format.space_before = Pt(4)
    codigo.paragraph_format.space_after = Pt(8)
    codigo.paragraph_format.line_spacing = 1.0
    codigo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def configurar_pagina_y_margenes(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def configurar_header_footer(doc: Document) -> None:
    section = doc.sections[0]

    header = section.header
    p_header = header.paragraphs[0]
    p_header.text = "AgroSmart — TFI Fundamentos de IA"
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p_header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_CAPTION

    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.add_run("Página ").font.size = Pt(9)
    _agregar_campo(p_footer, "PAGE", "1")
    p_footer.add_run(" de ").font.size = Pt(9)
    _agregar_campo(p_footer, "NUMPAGES", "1")
    for run in p_footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_CAPTION


# =====================================================================
# Helpers de alto nivel (semántica del informe)
# =====================================================================
def parrafos(doc: Document, *textos: str) -> None:
    for texto in textos:
        doc.add_paragraph(texto)


def lista_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def lista_numerada(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def codigo_bloque(doc: Document, lineas: str) -> None:
    """Agrega un bloque de código (cada línea = un párrafo de estilo CodigoMono)."""
    for linea in lineas.rstrip("\n").split("\n"):
        p = doc.add_paragraph(linea or " ", style="CodigoMono")
        # Sombreado claro para distinguir el bloque del texto normal.
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)


def figura(doc: Document, ruta: Path, numero: int, caption: str) -> None:
    """Embebe una imagen + caption numerado debajo."""
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ruta.exists():
        p_img.add_run().add_picture(str(ruta), width=ANCHO_FIGURA)
    else:
        run = p_img.add_run(f"[FIGURA NO ENCONTRADA: {ruta.name}]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(f"Figura {numero}: {caption}")
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION


def figura_placeholder(doc: Document, numero: int, caption: str, ancho_cm: float = 12.0, alto_cm: float = 7.0) -> None:
    """Marco vacío + caption. Para insertar la imagen manualmente en Word."""
    tabla = doc.add_table(rows=1, cols=1)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tabla.cell(0, 0)
    cell.width = Cm(ancho_cm)
    _set_cell_border(cell, top=True, bottom=True, left=True, right=True)
    _shade_cell(cell, "FAFAFA")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"\n\n\n[ INSERTAR AQUÍ MANUALMENTE EN WORD ]\n\n"
        f"Captura de pantalla de la aplicación Streamlit.\n\n\n"
    )
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(f"Figura {numero}: {caption}")
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION


def tabla_simple(doc: Document, headers: list[str], filas: list[list[str]], ancho_total_cm: float = 16.0) -> None:
    """Tabla con encabezado en negrita y bordes simples."""
    tabla = doc.add_table(rows=1, cols=len(headers))
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False

    ancho_col = Cm(ancho_total_cm / len(headers))
    fila_header = tabla.rows[0].cells
    for i, h in enumerate(headers):
        fila_header[i].text = ""
        p = fila_header[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        fila_header[i].width = ancho_col
        fila_header[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(fila_header[i], "D9E2F3")
        _set_cell_border(fila_header[i], top=True, bottom=True, left=True, right=True)

    for fila in filas:
        celdas = tabla.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = ""
            p = celdas[i].paragraphs[0]
            run = p.add_run(str(valor))
            run.font.size = Pt(10)
            celdas[i].width = ancho_col
            _set_cell_border(celdas[i], top=True, bottom=True, left=True, right=True)


# =====================================================================
# Carátula
# =====================================================================
def agregar_caratula(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titulo.add_run("AgroSmart")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITULO

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run(
        "Sistema híbrido de decisión agronómica para Argentina\n"
        "basado en Prolog y Random Forest"
    )
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = COLOR_SUBTITULO

    for _ in range(3):
        doc.add_paragraph()

    bloque = doc.add_paragraph()
    bloque.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bloque.add_run("Trabajo Final Integrador\n").font.size = Pt(14)
    bloque.add_run("Materia: Fundamentos de Inteligencia Artificial\n").font.size = Pt(12)
    bloque.add_run("Carrera: Maestría en Gestión y Desarrollo IA\n").font.size = Pt(12)
    bloque.add_run("Institución: Universidad CAECE\n").font.size = Pt(12)
    for run in bloque.runs:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    for _ in range(6):
        doc.add_paragraph()

    autores = doc.add_paragraph()
    autores.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = autores.add_run("Autores")
    run.font.bold = True
    run.font.size = Pt(12)
    autores2 = doc.add_paragraph()
    autores2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    autores2.add_run("Fernando Cáceres\nEzequiel Díaz Fernández").font.size = Pt(12)

    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run("\nMayo 2026").font.size = Pt(12)

    _saltar_pagina(doc)


# =====================================================================
# Resumen ejecutivo
# =====================================================================
def agregar_resumen_ejecutivo(doc: Document) -> None:
    doc.add_heading("Resumen ejecutivo", level=1)

    parrafos(
        doc,
        "AgroSmart es un sistema de recomendación de cultivos extensivos para Argentina "
        "que combina razonamiento simbólico (Prolog) con predicción cuantitativa "
        "(Random Forest). Dado un lote caracterizado por su región productiva y un "
        "conjunto reducido de variables agronómicas (suelo y clima), el sistema "
        "responde tres preguntas en cascada: qué cultivos son aptos, cuánto se estima "
        "producir y qué riesgos enfrenta el lote.",

        "El sistema se construyó a partir de un dataset maestro consolidado en "
        "tiempo real desde tres APIs públicas: el Ministerio de Agricultura, "
        "Ganadería y Pesca (MAGyP) para producción y rendimientos, Open-Meteo para "
        "clima histórico de reanálisis ERA5, y SoilGrids para propiedades físico-químicas "
        "del suelo. La consolidación cubre 11 cultivos extensivos sobre 30 departamentos "
        "representativos de las regiones Pampeana, NOA y NEA, con 25 campañas (2000/01 "
        "a 2024/25), totalizando 3.786 registros georreferenciados.",

        "El componente cuantitativo entrena un Random Forest independiente por cultivo "
        "sobre 12 features de suelo y clima. Los resultados van desde un R² de 0,82 para "
        "algodón —el cultivo mejor predicho, donde la capacidad de intercambio catiónico "
        "explica buena parte de la varianza— hasta un R² de 0,02 para maní, cuyo "
        "rendimiento depende fundamentalmente de variables que no están en el dataset. "
        "El componente simbólico está implementado en SWI-Prolog y aplica reglas "
        "de aptitud, riesgo y recomendación cuyos umbrales no se hardcodearon: provienen "
        "de percentiles calculados sobre el propio dataset argentino.",

        "El aporte propio del trabajo es la cascada que integra ambas capas. Para un "
        "lote dado, la consulta a Prolog identifica los cultivos aptos y los riesgos "
        "agronómicos; para cada cultivo aprobado, el modelo cuantitativo predice "
        "rendimiento e intervalo de confianza al 95% y el sistema confronta esa "
        "predicción contra los percentiles esperados de la zona, marcándola explícitamente "
        "si cae bajo el P10. La evaluación de un lote es instantánea (< 1 segundo) y "
        "se expone tanto vía CLI como vía una aplicación web Streamlit pensada para la "
        "defensa oral.",

        "El diferencial del sistema respecto de un clasificador puro es la explicabilidad: "
        "cada recomendación viene acompañada de la justificación simbólica que la sustenta, "
        "y cada predicción cuantitativa se interpreta dentro del rango histórico de la "
        "zona. El diferencial respecto de un sistema experto puro es la cuantificación: "
        "el productor no solo sabe qué cultivos son aptos, sino cuánto puede esperar "
        "producir y con qué incertidumbre.",
    )

    _saltar_pagina(doc)


# =====================================================================
# Índice (campo TOC, lo refresca Word)
# =====================================================================
def agregar_indice(doc: Document) -> None:
    doc.add_heading("Índice", level=1)

    nota = doc.add_paragraph()
    run = nota.add_run(
        "Este índice se genera automáticamente desde los títulos del documento. "
        "Si los números de página no aparecen al abrir el archivo en Microsoft Word, "
        "hacer click derecho sobre el índice y seleccionar “Actualizar campo” (o "
        "presionar F9 con el cursor sobre el índice)."
    )
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    p_toc = doc.add_paragraph()
    _agregar_campo(
        p_toc,
        'TOC \\o "1-3" \\h \\z \\u',
        "Click derecho aquí y elegir 'Actualizar campo' (o F9) para generar el índice.",
    )

    _saltar_pagina(doc)


# =====================================================================
# Capítulo 1 — Introducción
# =====================================================================
def agregar_capitulo_1(doc: Document) -> None:
    doc.add_heading("Capítulo 1 — Introducción y problemática", level=1)

    doc.add_heading("1.1 Contexto: la decisión agronómica en Argentina", level=2)
    parrafos(
        doc,
        "La agricultura extensiva argentina genera más del 50% de las exportaciones del "
        "país y opera en condiciones edafoclimáticas heterogéneas que requieren decisiones "
        "específicas por lote. Los principales cultivos —soja, maíz, trigo, girasol, "
        "sorgo y cereales menores— se siembran en cinco regiones productivas (Pampeana, "
        "NOA, NEA, Cuyo y Patagonia) con suelos, regímenes pluviométricos y temperaturas "
        "muy distintas. La elección del cultivo a sembrar en una campaña dada combina "
        "conocimiento agronómico de manual (rotaciones recomendadas, sensibilidad a "
        "heladas, requerimientos hídricos) con información cuantitativa local (textura "
        "del suelo, materia orgánica, lluvia esperada).",

        "En la práctica, esa decisión la toma habitualmente un asesor o ingeniero "
        "agrónomo a partir de su experiencia y de información parcial. Existen sistemas "
        "comerciales que asisten esa decisión, pero raramente combinan razonamiento "
        "simbólico explícito (las reglas que un agrónomo aplicaría) con modelos "
        "estadísticos calibrados sobre datos locales. El proyecto AgroSmart responde a "
        "esa carencia con un sistema híbrido cuyo razonamiento es trazable y cuyas "
        "predicciones cuantitativas se confrontan contra percentiles históricos para "
        "evitar que el modelo entregue salidas sin contexto.",
    )

    doc.add_heading("1.2 Por qué un sistema híbrido", level=2)
    parrafos(
        doc,
        "Un clasificador puramente estadístico, entrenado para responder “qué cultivo "
        "sembrar”, tiene una limitación estructural: opera como caja negra. Si predice "
        "que conviene maíz en lugar de soja para un lote dado, no entrega la "
        "justificación agronómica subyacente. Esto es problemático en un dominio donde "
        "el productor necesita comprender el “porqué” para tomar la decisión final, y "
        "donde un asesor debe poder defender su recomendación frente al cliente.",

        "Por otro lado, un sistema experto puramente simbólico —reglas Prolog basadas "
        "en literatura agronómica— responde el “porqué” pero no cuantifica. Puede decir "
        "“soja es apta acá”, pero no estima cuánto se va a producir ni con qué "
        "incertidumbre. Esa cuantificación es esencial para planificar márgenes "
        "económicos.",

        "El enfoque híbrido combina ambos paradigmas. La capa simbólica decide qué "
        "cultivos son aptos y diagnostica los riesgos. La capa estadística predice el "
        "rendimiento esperado para los cultivos aprobados. Y un puente bidireccional "
        "permite que las dos capas se contrasten mutuamente: si el modelo predice un "
        "rendimiento muy por debajo de los percentiles esperados de la zona, la "
        "advertencia se reporta explícitamente al usuario.",
    )

    doc.add_heading("1.3 Objetivos del trabajo", level=2)
    parrafos(
        doc,
        "El objetivo general fue diseñar e implementar un sistema híbrido de "
        "recomendación agronómica para Argentina que integre razonamiento simbólico "
        "y modelos cuantitativos sobre un dataset construido a partir de fuentes "
        "públicas oficiales.",

        "Los objetivos específicos fueron cinco. Primero, recuperar e integrar datos "
        "de producción, clima y suelo desde APIs argentinas e internacionales sin "
        "depender de datasets estáticos preexistentes. Segundo, derivar "
        "estadísticamente —desde el propio dataset— los rangos óptimos de cada cultivo "
        "que alimentan al sistema experto, evitando hardcodear umbrales arbitrarios. "
        "Tercero, construir un sistema experto Prolog modular que razone sobre aptitud, "
        "riesgo y recomendación. Cuarto, entrenar modelos cuantitativos por cultivo "
        "que predigan rendimiento e intervalo de confianza. Quinto, integrar ambas "
        "capas en una cascada de decisión exponible vía CLI y vía aplicación web.",
    )

    doc.add_heading("1.4 Alcance y limitaciones", level=2)
    parrafos(
        doc,
        "El sistema cubre 11 cultivos extensivos: soja, maíz, trigo, girasol, sorgo, "
        "cebada, avena, centeno, arroz, algodón y maní. La cobertura geográfica "
        "efectiva abarca tres regiones —Pampeana, NOA y NEA— sobre 30 departamentos "
        "representativos seleccionados a partir de la cobertura de MAGyP.",

        "Las regiones Cuyo y Patagonia quedaron sin cobertura efectiva: el "
        "Ministerio de Agricultura no reporta como cultivos extensivos los cereales "
        "que se siembran en Cuyo bajo riego (maíz, trigo, cebada en Mendoza) ni la "
        "avena para verdeo de Patagonia. Esta es una limitación de la fuente de "
        "datos, no del sistema. Los modelos podrían entrenarse para esas regiones "
        "si se incorporara una fuente alternativa de rendimientos.",

        "El alcance temporal del dataset es 2000/01–2024/25 (25 campañas). El corte "
        "en el año 2000 responde a una decisión metodológica documentada en la "
        "bitácora del proyecto: la introducción comercial de la soja transgénica "
        "(1996) y la consolidación de la siembra directa (1999-2001) "
        "reconfiguraron el paquete tecnológico del agro pampeano, y mezclar "
        "campañas anteriores con posteriores produciría datos no comparables.",
    )

    doc.add_heading("1.5 Estructura del informe", level=2)
    parrafos(
        doc,
        "El informe se organiza por componentes técnicos del sistema. El Capítulo 2 "
        "describe la recuperación e integración de datos desde APIs públicas y los "
        "problemas técnicos resueltos durante la consolidación. El Capítulo 3 "
        "presenta los hallazgos del análisis exploratorio y la derivación "
        "estadística de los rangos óptimos por cultivo. El Capítulo 4 documenta "
        "el sistema experto Prolog (hechos, reglas y API pública). El Capítulo 5 "
        "presenta los modelos cuantitativos Random Forest entrenados por cultivo "
        "y sus métricas. El Capítulo 6 detalla el bridge Python ↔ Prolog que "
        "integra ambas capas en una cascada de decisión. El Capítulo 7 describe "
        "la aplicación web Streamlit que envuelve el sistema para uso interactivo. "
        "El Capítulo 8 cierra con conclusiones, limitaciones reconocidas y trabajo "
        "futuro.",
    )

    _saltar_pagina(doc)


# =====================================================================
# Capítulo 2 — Recuperación e integración de datos
# =====================================================================
def agregar_capitulo_2(doc: Document) -> None:
    doc.add_heading("Capítulo 2 — Recuperación e integración de datos", level=1)

    doc.add_heading("2.1 Decisión de descartar datasets estáticos", level=2)
    parrafos(
        doc,
        "En la primera iteración del proyecto se contempló utilizar el dataset "
        "Crop Recommendation Dataset (Kaggle, autor Atharva Ingle) como fuente "
        "principal. Contiene N, P, K, temperatura, humedad, pH y lluvia para 22 "
        "cultivos. Sin embargo, el dataset es de origen indio: sus rangos óptimos "
        "están calibrados para suelos y climas indios, no argentinos. Como el "
        "sistema apunta específicamente a operadores agrícolas argentinos, "
        "trabajar con ese dataset introduciría sesgo sistemático en las reglas "
        "derivadas.",

        "La decisión que se tomó fue descartarlo y construir un dataset maestro "
        "en tiempo real desde tres APIs públicas. El trade-off: más complejidad "
        "de ingeniería (clientes HTTP, manejo de cuotas, cache local) a cambio "
        "de un dataset 100% argentino con trazabilidad completa de cada registro. "
        "Defensivamente, también es más sólido afirmar que las reglas del "
        "sistema experto se derivan de evidencia argentina real que de un "
        "dataset genérico de un contexto agronómico distinto.",
    )

    doc.add_heading("2.2 MAGyP: producción y rendimiento", level=2)
    parrafos(
        doc,
        "La fuente primaria es el portal de datos abiertos del Ministerio de "
        "Agricultura, Ganadería y Pesca de la Nación. Provee un único CSV "
        "maestro con todos los cultivos, todas las provincias y todos los "
        "departamentos desde 1969. El archivo pesa aproximadamente 15 MB "
        "y contiene 160.499 filas con esquema homogéneo: cultivo, año, "
        "campaña, provincia (con id INDEC), departamento, superficie sembrada, "
        "superficie cosechada, producción y rendimiento (kg/ha).",

        "Se eligió esta fuente única en lugar de los CSVs por cultivo "
        "individuales por tres razones: una sola descarga, mismo corte temporal "
        "entre cultivos y esquema homogéneo. La URL primaria contiene la fecha "
        "de actualización del corte (formato “2026-03”), que cambia "
        "periódicamente. Para garantizar que el cliente siga funcionando "
        "cuando esa URL quede obsoleta, se implementó un mecanismo de fallback "
        "que resuelve dinámicamente la URL vigente vía el endpoint CKAN del "
        "catálogo.",
    )

    doc.add_heading("2.3 Open-Meteo: clima histórico", level=2)
    parrafos(
        doc,
        "El servicio Open-Meteo provee acceso público al reanálisis ERA5 del "
        "European Centre for Medium-Range Weather Forecasts (ECMWF), con cobertura "
        "1940-presente y resolución espacial nativa de aproximadamente 9 km. Para "
        "cada combinación departamento-cultivo-campaña se solicita un resumen "
        "agroclimático del ciclo correspondiente: temperatura media, máxima y "
        "mínima, precipitación total, humedad relativa, radiación solar y días "
        "con helada.",

        "La lógica de agregación distingue entre cultivos de verano (soja, maíz, "
        "sorgo, girasol, arroz, algodón, maní) cuyo ciclo va de octubre a marzo, "
        "y cultivos de invierno (trigo, cebada, avena, centeno) cuyo ciclo va de "
        "abril a noviembre. Esta distinción es crítica porque la humedad y la "
        "temperatura relevantes para una soja de campaña difieren sustancialmente "
        "de las de un trigo del mismo año.",
    )

    doc.add_heading("2.4 SoilGrids: suelo georreferenciado", level=2)
    parrafos(
        doc,
        "El servicio SoilGrids del International Soil Reference and Information "
        "Centre (ISRIC) provee mapas globales de propiedades edáficas con "
        "resolución nativa de 250 m. Para cada coordenada se solicitan seis "
        "propiedades: arcilla, arena, limo, carbono orgánico (que se convierte a "
        "materia orgánica con el factor Van Bemmelen), pH y capacidad de "
        "intercambio catiónico (CEC). Las propiedades se piden a tres "
        "profundidades (0-5 cm, 5-15 cm, 15-30 cm) y se agregan a un valor "
        "único 0-30 cm con promedio ponderado por espesor, que es la zona "
        "radicular efectiva para los cultivos extensivos del proyecto.",
    )

    doc.add_heading("2.5 Problemas resueltos durante la consolidación", level=2)
    parrafos(
        doc,
        "La integración de tres servicios heterogéneos sobre 30 departamentos × "
        "11 cultivos × 25 campañas implicó resolver problemas reales de calidad "
        "de datos, cuotas de servicios públicos y desambiguación geográfica. Los "
        "principales se resumen a continuación; el detalle completo está "
        "documentado en la sección 4 de la bitácora del proyecto.",
    )

    doc.add_heading("Agujeros sistemáticos del SoilGrids en zonas urbanas", level=3)
    parrafos(
        doc,
        "Las primeras consultas a SoilGrids para los centroides de los "
        "departamentos cabecera (Pergamino, Río Cuarto, Junín) devolvieron valores "
        "nulos en todas las propiedades. Un barrido espacial sistemático demostró "
        "que la API excluye píxeles correspondientes a zonas urbanas y cuerpos de "
        "agua para no contaminar la grilla agronómica. Como muchos centroides "
        "departamentales coinciden con sus ciudades cabecera, la probabilidad de "
        "caer en estos huecos es alta.",

        "La solución implementada es un mecanismo de fallback espacial en dos "
        "anillos concéntricos. El primer anillo (radio aproximado de 1,1 km) "
        "consulta ocho vecinos cardinales y diagonales; el segundo (radio "
        "aproximado de 3,3 km) repite la consulta para los siguientes ocho "
        "vecinos. Solo se promedian los vecinos que devolvieron datos válidos. "
        "Las llamadas se ejecutan en paralelo para reducir la latencia. Sobre "
        "las 2.338 filas pampeanas iniciales, solo 134 quedaron sin dato (5,7%), "
        "todas correspondientes a un único departamento (General Pueyrredón, "
        "Mar del Plata) donde la urbanización costera produce un agujero más "
        "grande que el segundo anillo.",
    )

    doc.add_heading("Rate limiting de Open-Meteo (HTTP 429)", level=3)
    parrafos(
        doc,
        "Durante la primera corrida masiva de consolidación, Open-Meteo comenzó "
        "a devolver HTTP 429 (“Too Many Requests”) tras 11 minutos. El cliente "
        "original trataba todos los códigos 4xx como errores semánticos y los "
        "skipeaba inmediatamente, lo cual es correcto para 400 ó 404, pero no "
        "para 429 que es una señal de control de flujo, no de error.",

        "La solución se implementó en tres capas. Primero, un manejo específico "
        "del código 429 con backoff exponencial largo (300, 600 y 1200 segundos) "
        "y respeto del header Retry-After si la respuesta lo trae. Segundo, "
        "throttling preventivo con un sleep mínimo de 750 ms entre requests "
        "reales (no afecta cache hits), garantizando un máximo cercano a 80 "
        "calls por minuto, muy por debajo del límite de 600/min. Tercero, cache "
        "persistente en disco para que las regeneraciones del dataset no consuman "
        "cuota.",
    )

    doc.add_heading("Validación de nombres de departamentos", level=3)
    parrafos(
        doc,
        "Los nombres de departamentos en MAGyP siguen la nomenclatura oficial "
        "INDEC, con tildes y caracteres exactos. Cualquier desviación "
        "ortográfica produce que el filtro devuelva cero filas. Antes de "
        "aceptar la lista de 30 departamentos representativos, se validó cada "
        "nombre contra el CSV maestro descargado. Adicionalmente, se descubrió "
        "que cultivos como soja, trigo y cebada figuran en MAGyP con varias "
        "categorías (“soja 1ra”, “soja 2da”, “soja total”). Sumar las dos "
        "primeras produciría doble conteo; se decidió usar las versiones "
        "“total” en los tres casos.",
    )

    doc.add_heading("2.6 Dataset maestro consolidado", level=2)
    parrafos(
        doc,
        "El dataset final tiene 3.786 filas y 29 columnas. Las columnas se "
        "agrupan en cinco bloques: identificadoras (cultivo, región, provincia, "
        "departamento, campaña), geográficas (latitud, longitud), de MAGyP "
        "(superficies, producción, rendimiento), de Open-Meteo (temperatura "
        "media, máxima y mínima, precipitación, humedad relativa, radiación "
        "solar, días de helada) y de SoilGrids (arcilla, arena, limo, materia "
        "orgánica, pH, capacidad de intercambio catiónico). Tres columnas "
        "adicionales registran la calidad del dato de suelo: si vino directo o "
        "por qué anillo de fallback se rescató.",
    )

    tabla_simple(
        doc,
        ["Región", "Filas", "% del total", "Comentario"],
        [
            ["Pampeana", "2.672", "70,6%", "Núcleo del dataset"],
            ["NOA", "604", "16,0%", "Salta, Tucumán, Jujuy, Santiago del Estero"],
            ["NEA", "510", "13,5%", "Corrientes, Chaco, Entre Ríos norte, Santa Fe norte"],
            ["Cuyo", "0", "0%", "MAGyP no reporta cultivos extensivos en secano"],
            ["Patagonia", "0", "0%", "MAGyP no reporta avena para verdeo"],
        ],
    )

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run("Tabla 1: Cobertura del dataset maestro por región.")
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    parrafos(
        doc,
        "La calidad del dato de suelo se distribuye así: 12,4% de las filas "
        "obtuvo el dato directo del píxel exacto, 58,4% lo rescató del anillo "
        "de 1,1 km, 25,7% del anillo de 3,3 km, y 3,5% quedó sin dato. La "
        "distribución se grafica a continuación.",
    )

    figura(
        doc,
        FIGURAS_DIR / "09_distribucion_suelo_calidad.png",
        1,
        "Distribución de la calidad del dato de suelo en las 3.786 filas del "
        "dataset maestro. La mayoría de las filas requirió el mecanismo de "
        "fallback espacial implementado para resolver los agujeros del SoilGrids.",
    )

    _saltar_pagina(doc)


# =====================================================================
# Capítulo 3 — EDA y derivación de rangos óptimos
# =====================================================================
def agregar_capitulo_3(doc: Document) -> None:
    doc.add_heading("Capítulo 3 — Análisis exploratorio y derivación de rangos óptimos", level=1)

    doc.add_heading("3.1 Hallazgos clave del EDA", level=2)
    parrafos(
        doc,
        "El análisis exploratorio sobre el dataset consolidado arrojó tres tipos "
        "de hallazgos relevantes para el resto del proyecto. Primero, la "
        "heterogeneidad de rendimientos entre cultivos, que justifica entrenar "
        "un modelo separado por cultivo. Segundo, la captura efectiva de eventos "
        "climáticos extremos en el dataset, que valida la calidad del cruce "
        "entre las tres fuentes. Tercero, gradientes geográficos consistentes "
        "con la realidad agronómica argentina, que validan la coherencia del "
        "dato.",
    )

    figura(
        doc,
        FIGURAS_DIR / "01_boxplot_rendimiento_por_cultivo.png",
        2,
        "Boxplot del rendimiento (kg/ha) por cultivo. La heterogeneidad entre "
        "cultivos —tanto en mediana como en dispersión— es la primera "
        "justificación para entrenar modelos cuantitativos separados.",
    )

    doc.add_heading("Eventos climáticos extremos capturados", level=3)
    parrafos(
        doc,
        "El dataset captura con fidelidad eventos de sequía documentados como "
        "catastróficos en la región núcleo argentina. La campaña 2022/23 "
        "—la peor sequía en décadas, asociada al fenómeno La Niña triple— "
        "muestra en Pergamino solo 333,9 mm de lluvia (menos de la mitad del "
        "normal) y un rendimiento de soja de 1.857 kg/ha, casi 50% por debajo "
        "del histórico de la zona. Una campaña más antigua de sequía severa, "
        "2008/09, también queda capturada con 502 mm y 2.122 kg/ha. La "
        "correlación entre déficit hídrico y caída de rendimiento es "
        "claramente visible y es el tipo de señal que los modelos cuantitativos "
        "del Capítulo 5 aprovechan.",
    )

    figura(
        doc,
        FIGURAS_DIR / "04_evolucion_temporal_soja_pampeana.png",
        3,
        "Evolución temporal del rendimiento promedio de soja en la región "
        "pampeana, 2000/01 a 2024/25. Los eventos La Niña 2008/09 y 2022/23 "
        "quedan visibles como caídas pronunciadas que el dataset captura "
        "correctamente a partir del cruce entre rendimiento (MAGyP) y "
        "precipitación (Open-Meteo).",
    )

    doc.add_heading("Gradiente este-oeste pampeano", level=3)
    parrafos(
        doc,
        "La comparación entre Pergamino (Buenos Aires) y Río Cuarto (Córdoba) "
        "muestra el gradiente característico de la región pampeana: Pergamino "
        "tiene textura franco-limosa, materia orgánica alta (3,19%) y "
        "precipitación de 769 mm; Río Cuarto tiene textura franco-arenosa, "
        "materia orgánica menor (2,42%) y precipitación de 545 mm. El pH es "
        "ligeramente más alto en Río Cuarto (6,91 vs 6,53), consistente con "
        "un ambiente menos lixiviado por menor lluvia. Este gradiente es de "
        "manual de edafología argentina y aparece automáticamente en el dataset "
        "sin que nadie lo programe, lo cual es validación de calidad del cruce "
        "de las tres APIs.",
    )

    doc.add_heading("3.2 Derivación estadística de rangos óptimos", level=2)
    parrafos(
        doc,
        "Una decisión de diseño central del proyecto es que los rangos óptimos "
        "que alimentan al sistema experto Prolog no se hardcodean a partir de "
        "literatura agronómica genérica, sino que se derivan estadísticamente "
        "del propio dataset argentino. Para cada combinación cultivo-variable "
        "se calculan los percentiles 10 y 90 sobre las campañas con rendimiento "
        "por encima de la mediana del cultivo, que es el subconjunto que "
        "define qué condiciones acompañan a un buen resultado productivo.",

        "El criterio de filtrar por “rendimiento por encima de la mediana” "
        "—en lugar de simplemente “rendimiento positivo”— es importante. "
        "Calcular percentiles sobre la totalidad de las campañas mezclaría "
        "campañas catastróficas (que tiran los percentiles hacia los extremos) "
        "con campañas exitosas. La pregunta agronómicamente relevante no es "
        "“en qué condiciones se cultivó alguna vez”, sino “en qué condiciones "
        "se cultivó cuando el rendimiento fue alto”. El sub-dataset de "
        "campañas exitosas responde precisamente esa pregunta.",

        "Los percentiles se exportan automáticamente como hechos Prolog "
        "mediante un script Python (descripto en la sección 11 de la bitácora). "
        "Cada vez que el dataset se actualiza, basta con regenerar los hechos "
        "y el sistema experto pasa a usar los nuevos rangos sin que ninguna "
        "regla cambie.",
    )

    doc.add_heading("3.3 Tabla resumen de rangos óptimos", level=2)
    parrafos(
        doc,
        "A modo de muestra, la siguiente tabla presenta los percentiles 10 y "
        "90 de cinco variables agronómicas centrales para los cuatro cultivos "
        "principales del proyecto. Los valores provienen del archivo "
        "data/processed/rangos_optimos_por_cultivo.csv y son los que el "
        "sistema experto consume vía hechos rango_optimo/4.",
    )

    tabla_simple(
        doc,
        ["Cultivo", "Variable", "P10", "P90", "Mediana"],
        [
            ["Soja", "pH", "6,18", "6,96", "6,72"],
            ["Soja", "Precipitación (mm)", "478", "1.115", "714"],
            ["Soja", "Temperatura media (°C)", "20,9", "24,7", "22,5"],
            ["Soja", "Materia orgánica (%)", "2,42", "4,44", "2,69"],
            ["Soja", "Arcilla (%)", "18,5", "36,1", "26,5"],
            ["Maíz", "pH", "5,85", "6,96", "6,74"],
            ["Maíz", "Precipitación (mm)", "442", "972", "681"],
            ["Maíz", "Temperatura media (°C)", "20,7", "24,7", "22,3"],
            ["Maíz", "Materia orgánica (%)", "2,42", "4,94", "2,83"],
            ["Maíz", "Arcilla (%)", "18,5", "36,1", "26,5"],
            ["Trigo", "pH", "6,35", "6,96", "6,74"],
            ["Trigo", "Precipitación (mm)", "330", "728", "469"],
            ["Trigo", "Temperatura media (°C)", "12,8", "16,6", "14,9"],
            ["Trigo", "Materia orgánica (%)", "2,45", "4,94", "3,11"],
            ["Trigo", "Arcilla (%)", "18,5", "36,1", "27,0"],
            ["Girasol", "pH", "6,35", "6,96", "6,78"],
            ["Girasol", "Precipitación (mm)", "382", "849", "586"],
            ["Girasol", "Temperatura media (°C)", "20,6", "23,7", "21,8"],
            ["Girasol", "Materia orgánica (%)", "2,42", "4,69", "2,83"],
            ["Girasol", "Arcilla (%)", "18,5", "36,1", "26,5"],
        ],
    )

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(
        "Tabla 2: Rangos óptimos derivados estadísticamente para los cuatro "
        "cultivos principales (soja, maíz, trigo, girasol). Fuente: "
        "data/processed/rangos_optimos_por_cultivo.csv."
    )
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    figura(
        doc,
        FIGURAS_DIR / "13_heatmap_rangos_optimos_por_cultivo.png",
        4,
        "Heatmap completo de los rangos óptimos por cultivo y variable "
        "agronómica. Cada celda condensa P10, mediana y P90 de la variable "
        "para el cultivo correspondiente.",
    )

    doc.add_heading("3.4 Por qué los rangos no están hardcodeados", level=2)
    parrafos(
        doc,
        "La consigna del trabajo final exige aporte propio en Fase IV: "
        "integración real entre lo subsímbolico (datos) y lo simbólico (reglas). "
        "El primero de los dos puentes anunciados en el proyecto se materializa "
        "exactamente en este punto. Cualquier cuestionamiento sobre “de dónde "
        "salen esos umbrales” se responde mostrando la trazabilidad del "
        "pipeline: dataset_maestro.csv → notebook 02_eda_consolidado.ipynb → "
        "rangos_optimos_por_cultivo.csv → script generar_hechos_prolog.py → "
        "hechos_generados.pl. El sistema experto razona con evidencia argentina "
        "actualizable, no con umbrales fijos copiados de un manual.",
    )

    _saltar_pagina(doc)


# =====================================================================
# Capítulo 4 — Sistema experto Prolog
# =====================================================================
def agregar_capitulo_4(doc: Document) -> None:
    doc.add_heading("Capítulo 4 — Sistema experto Prolog", level=1)

    doc.add_heading("4.1 Arquitectura modular", level=2)
    parrafos(
        doc,
        "La capa simbólica del sistema está implementada en SWI-Prolog y se "
        "estructura en seis archivos dentro de src/prolog/. La separación "
        "responde a un criterio de responsabilidad: hay archivos que solo "
        "contienen hechos (datos), archivos que solo contienen reglas (lógica) "
        "y un archivo de entrada que carga todo en cascada.",
    )

    codigo_bloque(
        doc,
        "src/prolog/\n"
        "├── hechos_generados.pl      # auto-generado desde Python (no editar)\n"
        "├── hechos_expertos.pl       # conocimiento agronómico cargado a mano\n"
        "├── reglas_aptitud.pl        # ¿es apto este cultivo para este lote?\n"
        "├── reglas_riesgo.pl         # ¿qué riesgos enfrenta?\n"
        "├── reglas_recomendacion.pl  # ¿recomendar o no, con qué observaciones?\n"
        "└── agrosmart.pl             # punto de entrada (consulta de los 5 anteriores)\n"
    )

    parrafos(
        doc,
        "agrosmart.pl es el único archivo que Python necesita cargar. Mediante "
        "directivas :- consult/1 con paths relativos encadena todas las "
        "dependencias. SWI-Prolog resuelve los paths relativos al directorio del "
        "archivo que contiene la directiva, por lo que el sistema funciona con "
        "swipl src/prolog/agrosmart.pl desde cualquier directorio de trabajo.",
    )

    doc.add_heading("4.2 Hechos generados automáticamente", level=2)
    parrafos(
        doc,
        "El archivo hechos_generados.pl materializa el primer puente del aporte "
        "propio: datos → conocimiento. Lo produce un script Python "
        "(src/procesamiento/generar_hechos_prolog.py) que toma los percentiles "
        "calculados en el EDA y los exporta como hechos Prolog. El archivo "
        "contiene 265 hechos organizados en seis bloques.",
    )

    tabla_simple(
        doc,
        ["Predicado", "Aridad", "Hechos", "Origen"],
        [
            ["cultivo_soportado/1", "1", "11", "Lista canónica de cultivos del proyecto"],
            ["region_operacional/1", "1", "3", "Regiones con cobertura efectiva (pampeana, noa, nea)"],
            ["cultivo_en_region/2", "2", "20", "groupby (cultivo, region) sobre el dataset"],
            ["rango_optimo/4", "4", "110", "Percentiles 10 y 90 del CSV de rangos"],
            ["mediana_optima/3", "3", "110", "Mediana del CSV de rangos"],
            ["rendimiento_esperado/4", "4", "11", "Percentiles 10/50/90 del rendimiento"],
        ],
    )

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run("Tabla 3: Bloques de hechos generados automáticamente desde el dataset.")
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    parrafos(
        doc,
        "Después de regenerar el .pl, el propio script lo carga en SWI-Prolog "
        "vía pyswip y corre cinco consultas de aceptación. Si alguna falla, el "
        "script aborta con error explícito. Esta validación in-process garantiza "
        "que el archivo de hechos generado es siempre cargable: nunca se "
        "commitea un .pl roto al repositorio.",
    )

    doc.add_heading("4.3 Hechos expertos cargados a mano", level=2)
    parrafos(
        doc,
        "Hay conocimiento agronómico que no se puede derivar estadísticamente "
        "del dataset y debe codificarse a partir de bibliografía (INTA, FAO). "
        "Esto vive en hechos_expertos.pl. Incluye: el ciclo de cada cultivo "
        "(verano u invierno) que define qué datos climáticos son relevantes; la "
        "sensibilidad a heladas tempranas en otoño durante el llenado de grano "
        "para los siete cultivos de verano susceptibles; los tres cereales que "
        "necesitan vernalización (trigo, cebada, centeno); la categorización "
        "funcional por cultivo (principal, alternativo, forrajero, industrial); "
        "y nueve relaciones de predecesor recomendado para sostener rotaciones.",
    )

    doc.add_heading("4.4 Reglas de aptitud", level=2)
    parrafos(
        doc,
        "Las reglas de aptitud asumen que el lote bajo análisis tiene asociados "
        "dos predicados externos asertados por el bridge en tiempo de "
        "evaluación: region_lote(Lote, Region) y valor_lote(Lote, Variable, "
        "Valor) por cada variable agronómica. La capa simbólica está "
        "desacoplada: razona sobre el lote sin saber cómo se obtuvieron sus "
        "datos. Las reglas siguen una jerarquía de composición:",
    )

    codigo_bloque(
        doc,
        "% Una variable del lote cae dentro del rango óptimo del cultivo.\n"
        "en_rango_optimo(Lote, Cultivo, Variable) :-\n"
        "    valor_lote(Lote, Variable, Valor),\n"
        "    rango_optimo(Cultivo, Variable, Min, Max),\n"
        "    Valor >= Min,\n"
        "    Valor =< Max.\n"
        "\n"
        "% Aptitud edáfica: pH, materia orgánica y arcilla en rango.\n"
        "apto_suelo(Lote, Cultivo) :-\n"
        "    en_rango_optimo(Lote, Cultivo, ph),\n"
        "    en_rango_optimo(Lote, Cultivo, materia_organica_pct),\n"
        "    en_rango_optimo(Lote, Cultivo, arcilla_pct).\n"
        "\n"
        "% Aptitud climática: precipitación y temperatura media en rango.\n"
        "apto_clima(Lote, Cultivo) :-\n"
        "    en_rango_optimo(Lote, Cultivo, precipitacion_total_mm),\n"
        "    en_rango_optimo(Lote, Cultivo, temp_media_c).\n"
        "\n"
        "% Aptitud completa: viable geográficamente + suelo + clima.\n"
        "apto(Lote, Cultivo) :-\n"
        "    viable_geograficamente(Lote, Cultivo),\n"
        "    apto_suelo(Lote, Cultivo),\n"
        "    apto_clima(Lote, Cultivo).\n"
    )

    parrafos(
        doc,
        "Adicionalmente, dos predicados de aptitud parcial son útiles para no "
        "perder información. apto_parcial_suelo/2 captura los cultivos cuyo "
        "suelo es óptimo pero cuyo clima no lo es; apto_parcial_clima/2 hace "
        "lo simétrico. Muchos cultivos podrían ir bien en un lote con manejo "
        "específico (riego, encalado, variedad adaptada), y descartarlos por "
        "completo daría un reporte demasiado restrictivo.",
    )

    doc.add_heading("4.5 Reglas de riesgo", level=2)
    parrafos(
        doc,
        "El sistema diagnostica siete riesgos agronómicos. Tres de ellos "
        "—sequía, exceso hídrico y estrés térmico— usan umbrales cultivo-"
        "específicos derivados estadísticamente del dataset (los percentiles "
        "del archivo hechos_generados.pl). Los otros cuatro —helada, "
        "nutricional, acidez y alcalinidad— usan umbrales agronómicos "
        "absolutos válidos transversalmente para cualquier cultivo extensivo.",
    )

    tabla_simple(
        doc,
        ["Riesgo", "Condición", "Origen del umbral"],
        [
            ["Sequía", "Precipitación < P10 del cultivo", "Estadístico (dataset)"],
            ["Exceso hídrico", "Precipitación > P90 del cultivo", "Estadístico (dataset)"],
            ["Estrés térmico", "Temp. media > P90 del cultivo", "Estadístico (dataset)"],
            ["Helada", "Sensible_helada + días > 5", "Literatura INTA"],
            ["Nutricional", "MO < 2,0% y arena > 50%", "Literatura INTA"],
            ["Acidez", "pH < 5,5", "Literatura INTA"],
            ["Alcalinidad", "pH > 8,0", "Literatura INTA"],
        ],
    )

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run("Tabla 4: Riesgos diagnosticados por el sistema experto y su origen.")
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    doc.add_heading("4.6 Reglas de recomendación", level=2)
    parrafos(
        doc,
        "No todos los riesgos invalidan una recomendación: tres son críticos "
        "(sequía, helada, nutricional) y cuatro son no críticos (exceso "
        "hídrico, estrés térmico, acidez, alcalinidad). Los críticos descartan "
        "el cultivo porque no son fácilmente mitigables: la sequía es estrés "
        "hídrico de fondo, la helada es pérdida directa de área cosechable, el "
        "déficit nutricional es estructural del suelo. Los no críticos se "
        "reportan como observaciones para que el productor planifique las "
        "mitigaciones (drenaje, variedad de ciclo corto, encalado).",

        "El predicado público de más alto nivel es reporte_lote/4, que "
        "consume el bridge Python ↔ Prolog. Para un lote dado devuelve tres "
        "listas: cultivos recomendados, cultivos no recomendados y cultivos "
        "aptos parciales.",
    )

    codigo_bloque(
        doc,
        "% API pública del sistema experto.\n"
        "reporte_lote(Lote, Recomendados, NoRecomendados, AptosParciales) :-\n"
        "    findall(C, recomendar(Lote, C), Recomendados),\n"
        "    findall(C,\n"
        "            ( cultivo_soportado(C),\n"
        "              viable_geograficamente(Lote, C),\n"
        "              \\+ apto(Lote, C)\n"
        "            ),\n"
        "            NoRecomendados),\n"
        "    findall(C,\n"
        "            ( apto_parcial_suelo(Lote, C)\n"
        "            ; apto_parcial_clima(Lote, C)\n"
        "            ),\n"
        "            AptosParciales).\n"
    )

    doc.add_heading("4.7 Validación end-to-end", level=2)
    parrafos(
        doc,
        "Para validar el sistema experto se construyó un lote de prueba con "
        "datos típicos pampeanos (pH 6,5; materia orgánica 3,2%; arcilla 26%; "
        "precipitación 750 mm; temperatura media 22 °C; cero días de helada). "
        "El resultado fue exactamente el esperado por un agrónomo: cuatro "
        "cultivos recomendados (soja, maíz, sorgo, girasol —los grandes del "
        "verano pampeano), cinco no recomendados, dos aptos parciales (trigo "
        "y cebada, que son cereales de invierno cuyo suelo es perfecto para "
        "el lote pero cuyo clima no lo es porque los datos aportados "
        "corresponden a una campaña de verano).",

        "El sistema también descarta correctamente el arroz por "
        "viable_geograficamente/2: el hecho cultivo_en_region(arroz, pampeana) "
        "no existe porque el arroz solo se cultiva en NEA. Casos de borde "
        "como un lote árido (300 mm de lluvia) disparan el riesgo de sequía "
        "para soja por estar bajo el P10; un lote ácido (pH 5,0) dispara el "
        "riesgo de acidez. Las diez consultas de aceptación pasaron en una "
        "sola corrida del script de validación.",
    )

    _saltar_pagina(doc)


# =====================================================================
# Capítulo 5 — Modelos cuantitativos Random Forest
# =====================================================================
def agregar_capitulo_5(doc: Document) -> None:
    doc.add_heading("Capítulo 5 — Modelos cuantitativos (Random Forest)", level=1)

    doc.add_heading("5.1 Decisión: un modelo por cultivo", level=2)
    parrafos(
        doc,
        "La capa cuantitativa del sistema entrena un Random Forest "
        "independiente para cada uno de los 11 cultivos. Esta decisión —en "
        "lugar de un único modelo multitarea— se justifica agronómicamente "
        "desde el EDA: cada cultivo responde a drivers distintos. La "
        "capacidad de intercambio catiónico domina las predicciones de trigo "
        "y algodón; la humedad relativa domina soja y aparece en girasol; "
        "la radiación solar domina arroz y maní; la temperatura máxima es "
        "el driver principal de sorgo. Un único modelo para todos los "
        "cultivos disolvería estas diferencias y entregaría feature "
        "importances promediadas que no reflejan ninguna realidad agronómica "
        "concreta.",

        "La separación por cultivo permite, además, que el análisis de "
        "errores y las decisiones de mejora futuras se hagan sobre cada "
        "cultivo por separado, sin contaminación cruzada.",
    )

    figura(
        doc,
        FIGURAS_DIR / "07_correlaciones_top5_por_cultivo.png",
        5,
        "Top 5 correlaciones entre features agronómicas y rendimiento, por "
        "cultivo. La heterogeneidad entre cultivos justifica empíricamente "
        "la decisión de entrenar modelos separados.",
    )

    doc.add_heading("5.2 Hiperparámetros base sin GridSearch", level=2)
    parrafos(
        doc,
        "Los modelos se entrenan con configuración estándar de Random Forest, "
        "idéntica para los 11 cultivos: 200 estimadores, profundidad máxima "
        "15, mínimo 3 muestras por hoja, semilla 42 y paralelismo total de "
        "núcleos. La decisión consciente fue no usar GridSearch ni ningún "
        "otro mecanismo de fine-tuning. Las razones son tres.",
    )

    parrafos(
        doc,
        "Primero, el TFI no se evalúa por el R² del modelo: se evalúa por "
        "la integración entre la capa estadística y la simbólica. Optimizar "
        "0,05 puntos adicionales de R² no aporta valor académico al trabajo. "
        "Segundo, la interpretabilidad y simplicidad del pipeline son más "
        "importantes que el rendimiento marginal: cualquier evaluador puede "
        "leer cuatro hiperparámetros explícitos en el código, mientras que "
        "un GridSearchCV con 200 combinaciones sería ruido. Tercero, el "
        "tiempo de regeneración del sistema completo (entrenamiento de los "
        "11 modelos) es del orden de segundos con esta configuración. Esto "
        "es un objetivo de diseño: el TFI debe ser reproducible end-to-end "
        "en una computadora portátil sin GPU.",
    )

    doc.add_heading("5.3 Las 12 features en orden fijo", level=2)
    parrafos(
        doc,
        "El vector de entrada del modelo tiene 12 features en un orden "
        "documentado y fijo, definido como tupla constante en "
        "src/modelos/regresor_rendimiento.py. El orden es contrato entre el "
        "regresor y el bridge: el integrador construye el vector de "
        "inferencia desde los datos del lote en exactamente este orden antes "
        "de llamar a predict(). Cualquier cambio en el orden requiere "
        "reentrenar todos los modelos.",
    )

    codigo_bloque(
        doc,
        "FEATURES = (\n"
        "    'precipitacion_total_mm',\n"
        "    'temp_media_c',\n"
        "    'temp_max_promedio_c',\n"
        "    'temp_min_promedio_c',\n"
        "    'humedad_relativa_promedio',\n"
        "    'radiacion_solar_total',\n"
        "    'dias_helada',\n"
        "    'ph',\n"
        "    'materia_organica_pct',\n"
        "    'arcilla_pct',\n"
        "    'arena_pct',\n"
        "    'cec',\n"
        ")\n"
        "TARGET = 'rendimiento_kg_ha'\n"
    )

    doc.add_heading("5.4 Resultados por cultivo", level=2)
    parrafos(
        doc,
        "Tras entrenar los 11 modelos sobre las 3.523 filas válidas del "
        "dataset (filtrando rendimiento positivo y filas sin NaN en features), "
        "las métricas obtenidas se presentan en la tabla siguiente. Los datos "
        "exactos provienen del archivo data/modelos/reporte_entrenamiento.json "
        "generado durante el último entrenamiento.",
    )

    tabla_simple(
        doc,
        ["Cultivo", "n_train", "n_test", "R²", "MAE (kg/ha)", "RMSE (kg/ha)", "Top feature"],
        [
            ["Algodón", "92", "23", "0,818", "279", "330", "cec (0,65)"],
            ["Trigo", "406", "102", "0,747", "450", "563", "cec (0,56)"],
            ["Centeno", "89", "23", "0,741", "320", "406", "cec (0,22)"],
            ["Cebada", "86", "22", "0,710", "492", "593", "precip (0,31)"],
            ["Maíz", "532", "133", "0,638", "1.064", "1.368", "cec (0,26)"],
            ["Sorgo", "427", "107", "0,626", "687", "894", "temp_max (0,26)"],
            ["Soja", "507", "127", "0,577", "368", "479", "humedad (0,21)"],
            ["Avena", "232", "58", "0,480", "461", "594", "cec (0,17)"],
            ["Arroz", "56", "15", "0,437", "463", "667", "radiación (0,32)"],
            ["Girasol", "299", "75", "0,320", "328", "429", "humedad (0,14)"],
            ["Maní", "89", "23", "0,024", "610", "744", "radiación (0,22)"],
        ],
    )

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(
        "Tabla 5: Métricas de los modelos Random Forest por cultivo. "
        "Fuente: data/modelos/reporte_entrenamiento.json."
    )
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    doc.add_heading("5.5 Análisis de los extremos", level=2)
    parrafos(
        doc,
        "Algodón es el cultivo mejor predicho, con R² de 0,82, lo que significa "
        "que el modelo explica el 82% de la varianza del rendimiento de algodón "
        "en el dataset. La razón es estructural: el algodón en Argentina se "
        "concentra geográficamente en NOA y NEA con condiciones edafoclimáticas "
        "relativamente homogéneas, y la capacidad de intercambio catiónico "
        "domina la importancia con 0,65, indicando que un único valor edáfico "
        "explica buena parte de la varianza. Es un caso “fácil” en el sentido "
        "estadístico.",

        "Maní es prácticamente impredecible con las features actuales (R² de "
        "0,02). Esto no es un problema del modelo sino una limitación honesta "
        "del dataset: los drivers reales del rinde de maní son la variedad "
        "sembrada, la fecha de siembra exacta y la disponibilidad de calcio en "
        "el suelo, ninguna de las cuales está en el dataset. El sistema "
        "reconoce sus propios límites: para maní, la predicción cuantitativa "
        "entrega más ruido que señal y el R² lo evidencia. La consecuencia "
        "operativa es que el sistema solo recomendará maní cuando Prolog lo "
        "apruebe, pero la predicción asociada debe leerse con escepticismo.",
    )

    doc.add_heading("5.6 Intervalo de confianza al 95%", level=2)
    parrafos(
        doc,
        "Para cada predicción se entrega también un intervalo de confianza "
        "construido a partir de la dispersión de las predicciones de los "
        "árboles individuales del bosque. La lógica es la siguiente. Cada "
        "árbol del Random Forest se entrenó sobre un bootstrap distinto del "
        "dataset y por lo tanto es un estimador independiente. Para una "
        "entrada dada, los 200 árboles producen 200 predicciones individuales. "
        "La media de esas predicciones es la predicción del ensemble; la "
        "desviación estándar σ aproxima la incertidumbre del ensemble sobre "
        "ese punto. El intervalo del 95% se construye como [media − 1,96·σ, "
        "media + 1,96·σ].",

        "Este enfoque tiene tres ventajas: no requiere ningún supuesto "
        "distribucional sobre el ruido del target porque se deriva "
        "directamente de la varianza interna del bosque; es local porque la "
        "incertidumbre depende del punto evaluado (un lote en el corazón del "
        "espacio de entrenamiento tiene intervalo angosto, uno en la frontera "
        "tiene intervalo ancho); es interpretable porque se reporta en kg/ha "
        "junto con la predicción puntual.",
    )

    _saltar_pagina(doc)


# =====================================================================
# Capítulo 6 — Bridge Python ↔ Prolog
# =====================================================================
def agregar_capitulo_6(doc: Document) -> None:
    doc.add_heading("Capítulo 6 — Bridge Python ↔ Prolog", level=1)

    doc.add_heading("6.1 Justificación arquitectónica", level=2)
    parrafos(
        doc,
        "El bridge Python ↔ Prolog es el corazón del aporte propio del "
        "trabajo. Su función es orquestar la cascada de decisión que combina "
        "razonamiento simbólico (Prolog) con predicción cuantitativa "
        "(Random Forest) y producir un reporte integrado que el usuario "
        "consume. Sin este bridge, las dos capas existirían como sistemas "
        "paralelos sin diálogo; con él, se transforma en un sistema híbrido "
        "real donde cada capa potencia a la otra.",

        "La conexión se implementa con la librería pyswip, que expone la "
        "API de SWI-Prolog desde Python. El integrador encapsula esa "
        "dependencia detrás de una clase _BackendProlog que normaliza los "
        "tipos de retorno (átomos, listas) y permite intercambiar el "
        "backend en el futuro (por ejemplo, llamar a swipl como subprocess) "
        "sin afectar al resto del sistema.",
    )

    doc.add_heading("6.2 Las 3 etapas de la cascada", level=2)
    parrafos(
        doc,
        "El método principal evaluar_lote(lote) ejecuta tres etapas en "
        "secuencia, todas ocultas detrás de una API simple: el usuario pasa "
        "un objeto Lote y recibe un objeto ReporteLote serializable a JSON.",
    )

    doc.add_heading("Etapa 1 — Razonamiento simbólico", level=3)
    parrafos(
        doc,
        "Se asertan los hechos del lote en Prolog usando un identificador "
        "único (lote_runtime_<uuid>) para evitar contaminación entre "
        "llamadas concurrentes. Se consulta reporte_lote/4 para obtener las "
        "tres listas: recomendados, no recomendados y aptos parciales. "
        "Para cada cultivo recomendado se consulta también todos_los_riesgos/3 "
        "para enumerar las observaciones (riesgos no críticos). Al final, "
        "los hechos del lote se retractan, dejando la base de conocimiento "
        "limpia para la próxima evaluación.",
    )

    doc.add_heading("Etapa 2 — Predicción cuantitativa", level=3)
    parrafos(
        doc,
        "Para cada cultivo recomendado, el bridge carga el modelo joblib "
        "correspondiente desde data/modelos/. Construye el vector de features "
        "del lote en el orden documentado en el Capítulo 5 y predice "
        "rendimiento e intervalo de confianza al 95%. Si el modelo de un "
        "cultivo no estuviera entrenado (caso clonada limpia del repo sin "
        "haber corrido el entrenamiento), el bridge emite un warning y "
        "sigue funcionando con solo la parte simbólica.",
    )

    doc.add_heading("Etapa 3 — Reporte integrado", level=3)
    parrafos(
        doc,
        "Por cada cultivo recomendado se compara la predicción puntual "
        "contra los percentiles esperados de la zona (rendimiento_esperado/4 "
        "en Prolog) y se emite una clasificación cualitativa: alto, medio, "
        "bajo o muy_bajo. El reporte final es un objeto ReporteLote "
        "(dataclass) serializable a JSON, con: copia del lote, timestamp ISO, "
        "lista de recomendaciones (cada una con cultivo, predicción, "
        "intervalo, percentiles esperados, clasificación, observaciones), "
        "no recomendados con motivo, y aptos parciales con tipo.",
    )

    doc.add_heading("6.3 El segundo puente del aporte propio", level=2)
    parrafos(
        doc,
        "La cascada de la sub-sección anterior ya integra ambas capas. Pero "
        "el segundo puente del aporte propio se materializa en una decisión "
        "de diseño concreta dentro de la Etapa 3: si la predicción del "
        "modelo cuantitativo cae por debajo del P10 esperado de la zona, el "
        "bridge agrega automáticamente la observación "
        "rendimiento_bajo_lo_esperado a la lista de riesgos del cultivo, en "
        "lugar de aceptar la predicción en silencio.",

        "Esto es exactamente el “ML ↔ Lógica” anunciado en el planteo del "
        "trabajo. El sistema simbólico (que conoce los percentiles esperados "
        "de la zona) y el sistema estadístico (que predice un valor concreto "
        "para este lote) se contrastan mutuamente. Un desacuerdo no se "
        "oculta: se reporta como observación explícita para que el usuario "
        "comprenda que el modelo está prediciendo una campaña claramente por "
        "debajo de lo histórico.",

        "Es la diferencia entre dos sistemas paralelos y un sistema híbrido "
        "real. En un sistema paralelo, el modelo predice y Prolog opina por "
        "separado; en un sistema híbrido, las dos salidas se cruzan antes de "
        "presentar la recomendación.",
    )

    doc.add_heading("6.4 Casos de prueba: tres lotes contrastantes", level=2)
    parrafos(
        doc,
        "El script scripts/demo_agrosmart.py ejecuta tres lotes contrastantes "
        "que se usan para validar el sistema y para la defensa oral. La tabla "
        "siguiente resume los resultados.",
    )

    tabla_simple(
        doc,
        ["Lote", "Recomendados", "Observaciones"],
        [
            [
                "Pergamino típico",
                "soja, maíz, girasol, sorgo",
                "Las cuatro estimaciones quedan dentro del rango esperado de la zona pampeana.",
            ],
            [
                "Sequía 2022/23",
                "(ninguno)",
                "Todos los cultivos pampeanos disparan riesgo_sequia con 320 mm. Comportamiento agronómicamente correcto.",
            ],
            [
                "NEA arrocero",
                "arroz, maíz",
                "Predicción de arroz: 6.883 kg/ha (clasificación medio); soja descartada por pH 6,1 < P10 6,18.",
            ],
        ],
    )

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run("Tabla 6: Resultados de los tres lotes de demostración del sistema.")
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    doc.add_heading("6.5 Tests automatizados", level=2)
    parrafos(
        doc,
        "El bridge tiene cobertura mínima de tests automatizados en "
        "tests/test_bridge.py, con seis casos: un lote pampeano típico debe "
        "recomendar al menos soja y maíz; un lote árido (300 mm) debe "
        "excluir soja por riesgo de sequía; un lote ácido (pH 4,5) debe "
        "descartar todos los cultivos; la predicción de soja para un lote "
        "pampeano debe estar entre 1.500 y 4.500 kg/ha; el ReporteLote debe "
        "roundtrip-ear por JSON sin pérdida; el helper de clasificación "
        "cualitativa contra percentiles devuelve la categoría correcta. Los "
        "seis tests pasan en una corrida limpia (1,10 s).",
    )

    _saltar_pagina(doc)


# =====================================================================
# Capítulo 7 — App Streamlit
# =====================================================================
def agregar_capitulo_7(doc: Document) -> None:
    doc.add_heading("Capítulo 7 — Aplicación web (Streamlit)", level=1)

    doc.add_heading("7.1 Decisión de Streamlit", level=2)
    parrafos(
        doc,
        "Como complemento del CLI, se construyó una aplicación web que "
        "envuelve la cascada Python ↔ Prolog y la expone con una interfaz "
        "limpia para uso interactivo durante la defensa oral. La app no "
        "estaba en el alcance mínimo del TFI; es valor agregado pensado "
        "para que el evaluador pueda interactuar con el sistema en vivo, "
        "sin tener que leer código ni interpretar JSON crudo.",

        "El stack elegido fue Streamlit, en lugar de Flask con templates "
        "HTML o React con API REST. Las razones: Streamlit es Python puro "
        "(cero contexto nuevo para el desarrollador), los componentes son "
        "funciones que se renderizan a HTML automáticamente, los widgets "
        "manejan su propio estado en st.session_state, y no requiere "
        "custom CSS para verse profesional. La contracara —que cada "
        "interacción dispara un rerun completo del script— se resuelve con "
        "@st.cache_resource sobre la inicialización del SistemaAgroSmart.",
    )

    doc.add_heading("7.2 Componentes", level=2)
    parrafos(
        doc,
        "La app vive en el directorio app/ con un punto de entrada "
        "(streamlit_app.py) y cuatro componentes (sidebar, mapa, reporte y "
        "detalle por cultivo). El sidebar expone sliders para las siete "
        "variables del lote y un selector de “Cargar ejemplo” con cuatro "
        "lotes pre-configurados (Pergamino típico, Sequía 2022/23, NEA "
        "arrocero, Lote ácido) que reproducen los validados en el resto "
        "del proyecto. Esto evita tener que mover siete sliders en vivo "
        "durante la defensa: se elige el ejemplo y la cascada arranca con "
        "esos valores.",

        "El mapa interactivo renderiza los 30 departamentos del dataset "
        "como CircleMarker color-codeados por región (azul para pampeana, "
        "verde para NEA, naranja para NOA). Click en un marker autocompleta "
        "el sidebar con la mediana histórica del departamento.",

        "El reporte se organiza en tres tabs (Recomendados, No recomendados, "
        "Aptos parciales) reflejando la salida de reporte_lote/4. Cada "
        "cultivo recomendado aparece en una tarjeta con borde que contiene "
        "el banner de clasificación cualitativa (alto, medio, bajo o muy "
        "bajo), la predicción puntual con intervalo de confianza, la "
        "mediana histórica de la zona y los riesgos no críticos observados.",

        "El detalle por cultivo es la pieza clave del valor visual: en un "
        "expander, un selectbox lista todos los cultivos evaluados; al "
        "elegir uno se muestra una tabla “valor del lote vs rango óptimo” "
        "con check/cross visual variable por variable, y un boxplot Plotly "
        "horizontal con la distribución histórica del rendimiento del "
        "cultivo en la región del lote y la predicción del modelo marcada. "
        "El evaluador puede ver, variable por variable, exactamente por qué "
        "un cultivo es apto o no y dónde queda la predicción cuantitativa "
        "respecto del histórico.",
    )

    figura_placeholder(
        doc,
        6,
        "Captura de pantalla de la aplicación AgroSmart, evaluación del "
        "lote Pergamino típico.",
    )

    doc.add_heading("7.3 La app como herramienta de defensa", level=2)
    parrafos(
        doc,
        "La app permite al evaluador explorar interactivamente cualquier "
        "escenario en lugar de limitarse a los tres lotes pre-configurados. "
        "Mover un slider de precipitación de 750 a 320 mm muestra en vivo "
        "cómo todos los cultivos pampeanos pasan de “recomendados” a “no "
        "recomendados por riesgo_sequia”. Subir el pH de 6,5 a 4,5 muestra "
        "cómo el sistema descarta todos los cultivos por suelo fuera de "
        "rango. Esa interactividad es el principal valor visual del "
        "proyecto y es lo que hace que la defensa oral sea más persuasiva "
        "que la lectura de un JSON.",
    )

    _saltar_pagina(doc)


# =====================================================================
# Capítulo 8 — Conclusiones y trabajo futuro
# =====================================================================
def agregar_capitulo_8(doc: Document) -> None:
    doc.add_heading("Capítulo 8 — Conclusiones y trabajo futuro", level=1)

    doc.add_heading("8.1 Aportes del trabajo", level=2)
    parrafos(
        doc,
        "El trabajo entrega un sistema híbrido funcional que combina "
        "razonamiento simbólico explicable con predicción cuantitativa "
        "calibrada sobre datos argentinos reales. Tres aportes específicos "
        "merecen destacarse.",

        "Primero, los rangos óptimos que alimentan al sistema experto se "
        "derivan estadísticamente del propio dataset y no están hardcodeados. "
        "Esto significa que cualquier actualización de los datos (más "
        "campañas, nuevas regiones, diferente filtrado) se propaga "
        "automáticamente al razonamiento simbólico sin que ninguna regla "
        "Prolog cambie. Es la materialización concreta del primer puente "
        "del aporte propio: datos → conocimiento.",

        "Segundo, el bridge Python ↔ Prolog implementa una cascada de "
        "decisión donde la predicción cuantitativa se confronta contra los "
        "percentiles esperados de la zona y los desacuerdos se reportan "
        "explícitamente. No es un sistema donde dos componentes opinan en "
        "paralelo sino un sistema donde las dos capas se contrastan "
        "mutuamente. Es la materialización del segundo puente: ML ↔ Lógica.",

        "Tercero, el sistema es completamente reproducible end-to-end en "
        "una computadora portátil sin GPU. Una clonada limpia del repo, "
        "más cuatro comandos, regenera el dataset, los hechos Prolog, los "
        "modelos y la app. Esto es deliberado y es una propiedad valiosa "
        "para la trazabilidad académica.",
    )

    doc.add_heading("8.2 Limitaciones detectadas", level=2)
    parrafos(
        doc,
        "Ningún sistema es perfecto y este reconoce explícitamente sus "
        "limitaciones. Tres son particularmente honestas de mencionar.",

        "Cobertura geográfica incompleta. Las regiones Cuyo y Patagonia "
        "quedaron sin cobertura efectiva porque el Ministerio de "
        "Agricultura no reporta como cultivos extensivos los cereales "
        "que se siembran en Cuyo bajo riego, ni la avena para verdeo "
        "patagónica. El sistema podría operar sobre esas regiones si se "
        "incorporara una fuente alternativa, pero al día de hoy "
        "AgroSmart cubre efectivamente Pampeana, NOA y NEA, donde se "
        "concentra el grueso de la producción agrícola extensiva argentina.",

        "Maní impredecible con las features actuales. El R² de 0,02 para "
        "maní no es problema del modelo: los drivers reales del rinde de "
        "maní son la variedad genética, la fecha de siembra exacta y la "
        "disponibilidad de calcio en el suelo, ninguno de los cuales "
        "está en el dataset. El sistema reconoce sus propios límites.",

        "Comportamiento conservador en años catastróficos. Durante la "
        "sequía 2022/23, el sistema descarta absolutamente todos los "
        "cultivos por riesgo_sequia. Es agronómicamente correcto —un "
        "asesor que recomendara sembrar soja con 320 mm sería irresponsable— "
        "pero limita la utilidad práctica del sistema durante años "
        "catastróficos: un productor que igual va a sembrar merece una "
        "recomendación útil del tipo “si tenés que elegir, esto pierde "
        "menos”.",
    )

    figura(
        doc,
        FIGURAS_DIR / "11_campanias_fallidas_por_cultivo_y_campania.png",
        7,
        "Campañas con rendimiento cero o catastrófico por cultivo y año. "
        "Visualiza la concentración de fracasos en años de sequía severa "
        "(2008/09, 2022/23) y motiva la limitación documentada en 8.2 "
        "sobre el comportamiento conservador del sistema.",
    )

    doc.add_heading("8.3 Trabajo futuro", level=2)
    parrafos(
        doc,
        "Las extensiones naturales del sistema, en orden aproximado de "
        "prioridad, son las siguientes.",
    )

    lista_bullets(
        doc,
        [
            "Enriquecer el dataset con features adicionales que el modelo de "
            "maní necesita: variedad sembrada (vía registro INASE), fecha de "
            "siembra exacta y Growing Degree Days (GDD) calculados desde "
            "temperatura mínima y máxima diarias.",

            "Integrar disponibilidad de riego suplementario como variable "
            "del lote, lo cual permitiría recomendaciones útiles en "
            "regiones con sequía estacional.",

            "Construir un modelo de “pérdida proyectada” para campañas "
            "catastróficas, que rankee los cultivos descartados por "
            "magnitud de pérdida esperada en lugar de descartarlos en "
            "bloque.",

            "Desplegar la app web en un servidor público con autenticación "
            "para que asesores agronómicos puedan usarla en el campo.",

            "Expandir la cobertura a más cultivos no extensivos: papa, "
            "frutales, cultivos andinos. Implica resolver primero la "
            "cobertura de fuentes de datos en esas regiones.",

            "Persistir lotes evaluados en una base SQLite y exportar el "
            "reporte a PDF para compartir con el cliente final del asesor.",
        ],
    )

    _saltar_pagina(doc)


# =====================================================================
# Bibliografía
# =====================================================================
def agregar_bibliografia(doc: Document) -> None:
    doc.add_heading("Bibliografía y referencias", level=1)

    parrafos(
        doc,
        "El proyecto se apoyó fundamentalmente en documentación técnica "
        "oficial de las fuentes de datos y de las librerías utilizadas. "
        "Las referencias agronómicas son materiales de cátedra del INTA "
        "y guías de la FAO ampliamente difundidas.",
    )

    doc.add_heading("Fuentes de datos", level=2)
    lista_bullets(
        doc,
        [
            "Ministerio de Agricultura, Ganadería y Pesca de la Nación "
            "Argentina (MAGyP). Datos abiertos de estimaciones agrícolas. "
            "https://datos.magyp.gob.ar",

            "Open-Meteo. Historical Weather API (reanálisis ERA5). "
            "https://open-meteo.com/en/docs/historical-weather-api",

            "ISRIC – World Soil Information. SoilGrids 2.0. "
            "https://www.isric.org/explore/soilgrids",

            "European Centre for Medium-Range Weather Forecasts (ECMWF). "
            "ERA5 reanalysis. https://www.ecmwf.int/en/forecasts/dataset/ecmwf-reanalysis-v5",
        ],
    )

    doc.add_heading("Documentación técnica", level=2)
    lista_bullets(
        doc,
        [
            "SWI-Prolog Reference Manual. https://www.swi-prolog.org/pldoc/",
            "scikit-learn User Guide. https://scikit-learn.org/stable/user_guide.html",
            "Streamlit documentation. https://docs.streamlit.io/",
            "pyswip — Python interface to SWI-Prolog. https://pypi.org/project/pyswip/",
            "Folium — Python library for Leaflet.js maps. https://python-visualization.github.io/folium/",
            "Plotly Python. https://plotly.com/python/",
            "python-docx documentation. https://python-docx.readthedocs.io/",
        ],
    )

    doc.add_heading("Referencias agronómicas", level=2)
    lista_bullets(
        doc,
        [
            "Instituto Nacional de Tecnología Agropecuaria (INTA). Buenas "
            "prácticas agrícolas para la Argentina. Publicaciones técnicas "
            "varias. https://www.argentina.gob.ar/inta",

            "Food and Agriculture Organization (FAO). Crop Water "
            "Requirements (Irrigation and Drainage Paper No. 56, "
            "Penman-Monteith). https://www.fao.org/3/x0490e/x0490e00.htm",

            "Bolsa de Cereales de Buenos Aires. Panorama agrícola "
            "semanal (PAS) y reportes de campañas. https://www.bolsadecereales.com",
        ],
    )

    _saltar_pagina(doc)


# =====================================================================
# Anexo A — Repositorio y artefactos
# =====================================================================
def agregar_anexo_a(doc: Document) -> None:
    doc.add_heading("Anexo A — Repositorio y artefactos", level=1)

    doc.add_heading("A.1 URL del repositorio", level=2)
    parrafos(
        doc,
        "El código fuente completo del proyecto, junto con la bitácora "
        "extendida (docs/recuperacion_de_datos.md) y los scripts de "
        "regeneración, vive en un repositorio Git versionado. La URL "
        "exacta y las credenciales para acceder se proporcionan junto "
        "con la entrega formal del informe.",
    )

    doc.add_heading("A.2 Estructura del repositorio", level=2)
    codigo_bloque(
        doc,
        "agrosmart/\n"
        "├── data/\n"
        "│   ├── processed/              # dataset maestro y rangos por cultivo\n"
        "│   ├── modelos/                # joblib y reporte JSON (no versionado)\n"
        "│   └── cache/                  # cache de respuestas de APIs (no versionado)\n"
        "├── notebooks/\n"
        "│   └── 02_eda_consolidado.ipynb\n"
        "├── src/\n"
        "│   ├── apis/                   # clientes MAGyP, Open-Meteo, SoilGrids\n"
        "│   ├── procesamiento/          # consolidación + generación de hechos Prolog\n"
        "│   ├── modelos/                # regresor Random Forest por cultivo\n"
        "│   ├── prolog/                 # 6 archivos del sistema experto\n"
        "│   └── bridge/                 # cascada Python ↔ Prolog\n"
        "├── app/                        # aplicación Streamlit\n"
        "│   ├── streamlit_app.py\n"
        "│   └── componentes/            # sidebar, mapa, reporte, detalle\n"
        "├── scripts/\n"
        "│   ├── demo_agrosmart.py\n"
        "│   └── generar_informe.py      # genera este documento\n"
        "├── tests/\n"
        "│   └── test_bridge.py\n"
        "├── docs/\n"
        "│   ├── recuperacion_de_datos.md (bitácora extendida)\n"
        "│   ├── informe.docx             (este informe)\n"
        "│   └── figuras/                 # 13 PNGs del EDA\n"
        "├── requirements.txt\n"
        "├── CLAUDE.md\n"
        "└── README.md\n"
    )

    doc.add_heading("A.3 Instrucciones de reproducción", level=2)
    parrafos(
        doc,
        "Para regenerar el sistema completo desde una clonada limpia del "
        "repositorio, ejecutar los siguientes comandos en orden, parado en "
        "la raíz del proyecto:",
    )
    lista_numerada(
        doc,
        [
            "python -m venv .venv && .venv\\Scripts\\activate "
            "(en Windows) o source .venv/bin/activate (en Linux/Mac).",

            "pip install -r requirements.txt.",

            "Ejecutar el notebook notebooks/02_eda_consolidado.ipynb hasta "
            "el final, lo cual regenera "
            "data/processed/rangos_optimos_por_cultivo.csv y las 13 "
            "figuras de docs/figuras/.",

            "python -m src.procesamiento.generar_hechos_prolog "
            "(genera src/prolog/hechos_generados.pl).",

            "python -m src.modelos.regresor_rendimiento --entrenar "
            "(entrena los 11 modelos y genera "
            "data/modelos/reporte_entrenamiento.json).",

            "python scripts/demo_agrosmart.py "
            "(ejecuta los tres lotes de demostración del Capítulo 6).",

            "streamlit run app/streamlit_app.py "
            "(lanza la aplicación web descripta en el Capítulo 7).",

            "python scripts/generar_informe.py "
            "(regenera este informe en docs/informe.docx).",
        ],
    )

    doc.add_heading("A.4 Tests automatizados", level=2)
    parrafos(
        doc,
        "El proyecto incluye un conjunto mínimo de tests con pytest. Para "
        "correrlos: pytest tests/ -v. Los seis tests de tests/test_bridge.py "
        "validan el comportamiento end-to-end de la cascada (recomendaciones "
        "esperadas para lotes pampeanos típicos, descarte correcto en lotes "
        "áridos y ácidos, predicciones dentro de rango razonable, "
        "serialización JSON sin pérdida y clasificación cualitativa).",
    )

    doc.add_heading("A.5 Datos de contacto", level=2)
    parrafos(
        doc,
        "Para consultas sobre el proyecto:",
    )
    lista_bullets(
        doc,
        [
            "Fernando Cáceres — fcaceres78@gmail.com",
            "Ezequiel Díaz Fernández — (correo a completar)",
            "Universidad CAECE — https://www.caece.edu.ar",
        ],
    )


# =====================================================================
# Main
# =====================================================================
def main() -> int:
    print(f"Generando informe en {INFORME_PATH}...")
    doc = Document()
    configurar_estilos(doc)
    configurar_pagina_y_margenes(doc)
    configurar_header_footer(doc)

    agregar_caratula(doc)
    agregar_resumen_ejecutivo(doc)
    agregar_indice(doc)
    agregar_capitulo_1(doc)
    agregar_capitulo_2(doc)
    agregar_capitulo_3(doc)
    agregar_capitulo_4(doc)
    agregar_capitulo_5(doc)
    agregar_capitulo_6(doc)
    agregar_capitulo_7(doc)
    agregar_capitulo_8(doc)
    agregar_bibliografia(doc)
    agregar_anexo_a(doc)

    INFORME_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(INFORME_PATH)
    print(f"OK. Informe guardado en {INFORME_PATH}")
    print(f"  Tamano: {INFORME_PATH.stat().st_size / 1024:.1f} KB")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
