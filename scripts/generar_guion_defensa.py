"""
Genera el guion completo de defensa oral del TFI AgroSmart en
docs/AgroSmart_Guion_Defensa_TFI.docx.

Script reproducible: el contenido completo del guion está en este
archivo. Para regenerar: `python scripts/generar_guion_defensa.py`.

El documento es material de consulta bajo presión durante la defensa,
no un documento de densidad máxima: prioriza espaciado generoso,
tipografía clara y secciones inequívocas. Sigue el mismo patrón de
estilos que scripts/generar_informe.py para mantener consistencia
visual con el informe TFI.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------
# Rutas y constantes de estilo
# ---------------------------------------------------------------------
RAIZ = Path(__file__).resolve().parents[1]
DOCS_DIR = RAIZ / "docs"
SALIDA = DOCS_DIR / "AgroSmart_Guion_Defensa_TFI.docx"

FUENTE_CUERPO = "Calibri"
FUENTE_CODIGO = "Consolas"

COLOR_TITULO = RGBColor(0x1F, 0x4E, 0x79)        # azul oscuro académico
COLOR_SUBTITULO = RGBColor(0x2E, 0x75, 0xB6)     # azul medio
COLOR_BLOQUE_EZE = RGBColor(0x70, 0x30, 0xA0)    # violeta para bloques de Ezequiel
COLOR_FRASE_ANCLA = RGBColor(0xC0, 0x39, 0x2B)   # rojo discreto para frases literales
COLOR_CAPTION = RGBColor(0x59, 0x59, 0x59)
COLOR_FONDO_CODIGO = "F2F2F2"


# =====================================================================
# Helpers de bajo nivel
# =====================================================================
def _set_cell_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "808080")
        borders.append(border)
    tc_pr.append(borders)


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _agregar_campo(parrafo, instruccion: str, placeholder: str) -> None:
    """Inserta un campo OOXML (PAGE, NUMPAGES). Word lo evalúa al abrir."""
    run = parrafo.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruccion
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder_el = OxmlElement("w:t")
    placeholder_el.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, placeholder_el, fld_end):
        run._r.append(el)


def _saltar_pagina(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# =====================================================================
# Estilos globales
# =====================================================================
def configurar_estilos(doc: Document) -> None:
    estilos = doc.styles

    normal = estilos["Normal"]
    normal.font.name = FUENTE_CUERPO
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = 1.20
    pf.space_after = Pt(6)

    h1 = estilos["Heading 1"]
    h1.font.name = FUENTE_CUERPO
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_TITULO
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = estilos["Heading 2"]
    h2.font.name = FUENTE_CUERPO
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_SUBTITULO
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = estilos["Heading 3"]
    h3.font.name = FUENTE_CUERPO
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_TITULO
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


def configurar_pagina_y_footer(doc: Document) -> None:
    """Márgenes generosos + footer con paginación."""
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Página ")
    run.font.name = FUENTE_CUERPO
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_CAPTION
    _agregar_campo(p, "PAGE", "1")
    run = p.add_run(" de ")
    run.font.name = FUENTE_CUERPO
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_CAPTION
    _agregar_campo(p, "NUMPAGES", "1")


# =====================================================================
# Helpers de contenido (alto nivel)
# =====================================================================
def add_heading(doc: Document, texto: str, nivel: int = 1) -> None:
    doc.add_heading(texto, level=nivel)


def add_heading_eze(doc: Document, texto: str, nivel: int = 1) -> None:
    """Heading marcado visualmente como bloque de Ezequiel (violeta + tag)."""
    h = doc.add_heading(level=nivel)
    run = h.add_run("[EZE]  ")
    run.font.color.rgb = COLOR_BLOQUE_EZE
    run.font.bold = True
    run = h.add_run(texto)
    run.font.color.rgb = COLOR_BLOQUE_EZE


def add_parrafo(doc: Document, texto: str, *, justificado: bool = True) -> None:
    p = doc.add_paragraph(texto)
    p.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justificado else WD_ALIGN_PARAGRAPH.LEFT
    )


def add_caption(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION


def add_mensaje_clave(doc: Document, texto: str) -> None:
    """Caja destacada con la frase ancla del bloque, en negrita y color."""
    tabla = doc.add_table(rows=1, cols=1)
    tabla.autofit = True
    cell = tabla.rows[0].cells[0]
    _shade_cell(cell, "FFF2CC")  # amarillo suave
    _set_cell_border(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    run_lbl = p.add_run("MENSAJE CLAVE  ")
    run_lbl.bold = True
    run_lbl.font.size = Pt(9)
    run_lbl.font.color.rgb = COLOR_CAPTION
    p.add_run("\n")
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TITULO
    doc.add_paragraph()


def add_slide(doc: Document, etiqueta: str, contenido: str) -> None:
    """Bloque de código que reproduce literalmente el contenido del slide."""
    add_caption(doc, etiqueta)
    tabla = doc.add_table(rows=1, cols=1)
    tabla.autofit = True
    cell = tabla.rows[0].cells[0]
    _shade_cell(cell, COLOR_FONDO_CODIGO)
    _set_cell_border(cell)
    cell.text = ""
    # Cada línea como párrafo separado para preservar layout
    primero = True
    for linea in contenido.splitlines():
        p = cell.paragraphs[0] if primero else cell.add_paragraph()
        primero = False
        run = p.add_run(linea or " ")
        run.font.name = FUENTE_CODIGO
        run.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)


def add_bullets_eze(doc: Document, items: list[str]) -> None:
    """Bullets de Ezequiel marcados con prefijo violeta."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run("[EZE] ")
        run.bold = True
        run.font.color.rgb = COLOR_BLOQUE_EZE
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def add_frase_ancla(doc: Document, etiqueta: str, frase: str) -> None:
    """Frase literal en cajón rojo discreto. Para transiciones y cierres."""
    tabla = doc.add_table(rows=1, cols=1)
    cell = tabla.rows[0].cells[0]
    _shade_cell(cell, "FCEFEC")
    _set_cell_border(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(etiqueta + "\n")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_CAPTION
    p2 = cell.add_paragraph()
    run = p2.add_run(f'"{frase}"')
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_FRASE_ANCLA
    doc.add_paragraph()


def add_tips(doc: Document, titulo: str, items: list[str]) -> None:
    add_heading(doc, titulo, nivel=3)
    add_bullets(doc, items)


def add_tabla_simple(
    doc: Document,
    encabezados: list[str],
    filas: list[list[str]],
) -> None:
    """Tabla con encabezado sombreado + filas. Para plan B y checklist."""
    tabla = doc.add_table(rows=1 + len(filas), cols=len(encabezados))
    tabla.style = "Light Grid Accent 1"
    for i, h in enumerate(encabezados):
        cell = tabla.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        _shade_cell(cell, "D9E2F3")
    for j, fila in enumerate(filas):
        for i, valor in enumerate(fila):
            cell = tabla.rows[1 + j].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(valor)
            run.font.size = Pt(10)
    doc.add_paragraph()


# =====================================================================
# PORTADA
# =====================================================================
def render_portada(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guion de presentación oral")
    run.font.name = FUENTE_CUERPO
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITULO

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TFI AgroSmart")
    run.font.name = FUENTE_CUERPO
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_SUBTITULO

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Trabajo Final Integrador · Fundamentos de IA")
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = COLOR_CAPTION

    for _ in range(3):
        doc.add_paragraph()

    datos = [
        ("Universidad", "CAECE — Modalidad E-distancia"),
        ("Carrera", "Maestría en Gestión y Desarrollo de IA"),
        ("Materia", "Fundamentos de Inteligencia Artificial"),
        ("Profesor", "Juan Miguel Azcurra"),
        ("Cuatrimestre", "1° Cuatrimestre 2026"),
        ("Fecha de entrega", "22 de mayo de 2026"),
        ("Autores", "Fernando Cáceres · Ezequiel Díaz Fernández"),
    ]
    tabla = doc.add_table(rows=len(datos), cols=2)
    tabla.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (clave, valor) in enumerate(datos):
        c1 = tabla.rows[i].cells[0]
        c2 = tabla.rows[i].cells[1]
        c1.text = ""
        c2.text = ""
        p = c1.paragraphs[0]
        run = p.add_run(clave)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TITULO
        p = c2.paragraphs[0]
        run = p.add_run(valor)
        run.font.size = Pt(11)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Defensa oral · 30 min · Zoom")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_CAPTION

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Repositorio: https://github.com/fcaceres-create/caece-agro")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("App live: https://caece-agrosmart.streamlit.app")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_CAPTION

    _saltar_pagina(doc)


# =====================================================================
# INTRODUCCIÓN
# =====================================================================
def render_introduccion(doc: Document) -> None:
    add_heading(doc, "1. Introducción — estructura general", nivel=1)
    add_parrafo(
        doc,
        "Defensa de 30 minutos por Zoom con Prof. Juan Miguel Azcurra. "
        "Modalidad: presentación oral con slides + demo en vivo de la app. "
        "División de partes: Fernando ~20 min (apertura, datos, modelos, "
        "bridge, demo) + Ezequiel ~10 min (sistema experto Prolog, cierre)."
    )

    add_heading(doc, "Mapa de los 7 bloques (30 min)", nivel=2)
    encabezados = ["Tiempo", "Bloque", "Quién", "Slides / Escenas"]
    filas = [
        ["00–03", "Apertura", "Fernando", "3 slides"],
        ["03–07", "Datos y EDA", "Fernando", "3 slides"],
        ["07–13", "Sistema experto Prolog ⭐", "Ezequiel", "4 slides"],
        ["13–17", "Modelos Random Forest", "Fernando", "3 slides"],
        ["17–22", "Bridge en cascada ⭐⭐", "Fernando", "3 slides"],
        ["22–28", "Demo en vivo ⭐⭐⭐", "Fernando", "4 escenas"],
        ["28–30", "Cierre", "Ezequiel", "2 slides"],
        ["30+", "Q&A — ambos disponibles", "Ambos", "—"],
    ]
    add_tabla_simple(doc, encabezados, filas)
    add_caption(
        doc,
        "⭐ = momentos donde el profesor evaluará profundidad técnica."
    )

    add_heading(doc, "3 transiciones críticas (literales)", nivel=2)
    add_frase_ancla(
        doc,
        "Fernando → Ezequiel · final del bloque 2",
        "Sobre estos datos construimos las dos capas del sistema híbrido. "
        "Voy a dejarle la palabra a Ezequiel para que les cuente la primera, "
        "que es el sistema experto en Prolog."
    )
    add_frase_ancla(
        doc,
        "Ezequiel → Fernando · final del bloque 3",
        "Bueno, esa es la capa simbólica. Le devuelvo la palabra a Fer para "
        "que cuente la capa cuantitativa: los modelos Random Forest."
    )
    add_frase_ancla(
        doc,
        "Fernando → Ezequiel · final del bloque 6",
        "Esto es lo que construimos y cómo funciona. Le voy a dejar la "
        "palabra a Eze para que cierre con limitaciones y trabajo futuro."
    )

    _saltar_pagina(doc)


# =====================================================================
# BLOQUE 1 — APERTURA
# =====================================================================
def render_bloque_1(doc: Document) -> None:
    add_heading(doc, "2. Bloque 1 · Apertura  (3 min · Fernando)", nivel=1)
    add_mensaje_clave(
        doc,
        "Construimos un sistema híbrido que combina las dos paradigmas "
        "centrales de IA aplicado a un problema real argentino."
    )

    add_heading(doc, "Slide 1.1 — Portada  (~20 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
AGROSMART
Sistema híbrido de decisión agronómica para Argentina

Trabajo Final Integrador
Fundamentos de Inteligencia Artificial

Fernando Cáceres · Ezequiel Díaz Fernández
Prof. Juan Miguel Azcurra · 22/05/2026
Maestría en Gestión y Desarrollo de IA · CAECE""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "Buen día, profesor. Soy Fernando Cáceres. Junto a Ezequiel Díaz "
        "Fernández presentamos AgroSmart, nuestro Trabajo Final Integrador "
        "para Fundamentos de IA.",
        "Pausa de respiración. No te apures. Mirá a la cámara, no a las slides.",
    ])

    add_heading(doc, "Slide 1.2 — El problema  (~70 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
EL PROBLEMA

Un productor argentino tiene un lote.
Tres preguntas que no son triviales:

  1. ¿Qué cultivos son aptos para sembrar?
  2. ¿Cuánto va a producir?
  3. ¿Qué riesgos enfrenta?

Las respuestas existen — pero están dispersas:
  · Datos productivos en MAGyP
  · Datos climáticos en Open-Meteo
  · Datos de suelo en SoilGrids
  · Conocimiento agronómico en INTA, FAO, cátedra""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "El productor tiene fragmentos de la respuesta dispersos en distintas "
        "fuentes oficiales.",
        "MAGyP tiene 25 campañas reales de producción por departamento.",
        "Open-Meteo tiene reanálisis climático histórico.",
        "SoilGrids tiene propiedades del suelo georreferenciadas.",
        "Los expertos del INTA y la FAO tienen conocimiento agronómico "
        "estructurado.",
        "Pero nadie los integró en un sistema que tome decisiones. Esa es la "
        "oportunidad.",
    ])

    add_heading(doc, "Slide 1.3 — La tesis  (~70 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
NUESTRA TESIS

Un solo paradigma no alcanza:

  · Solo Random Forest → opaco
    "Soja con probabilidad 0.87. ¿Por qué? No sé."

  · Solo sistema experto → no cuantifica
    "Soja es apta. ¿Cuánto produce? No puede decir."

→ Combinarlos da una decisión simultáneamente
   EXPLICABLE y CUANTIFICABLE.

Las dos paradigmas centrales de la materia,
trabajando juntos.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "Si usábamos solo Random Forest, íbamos a tener una caja negra.",
        "Si usábamos solo sistema experto, podríamos razonar sobre aptitud y "
        "riesgo, pero no cuantificar cuánto va a rendir el lote.",
        "Nuestra tesis es que combinarlos da algo que ninguno puede dar solo: "
        "una decisión simultáneamente explicable (las reglas son auditables) "
        "y cuantificable (los modelos predicen).",
        "Y que eso es además honesto a la materia que estamos cursando: las "
        "dos paradigmas centrales de la IA clásica, trabajando juntos sobre "
        "un problema real.",
    ])

    add_tips(doc, "Tips del bloque 1", [
        "Hablá despacio. Los primeros 3 minutos siempre se aceleran por "
        "nervios. Conscientemente bajá el ritmo.",
        "No leas las slides. Las slides son referencia para el profesor; vos "
        "hablás de los conceptos.",
        "No menciones a Eze hasta el bloque 3. La apertura es tuya — clara "
        "y firme.",
        "Si el profesor te interrumpe con preguntas: respondé brevemente y "
        "volvé a las slides con \"justo el siguiente punto cubre eso\" o "
        "\"ese punto lo desarrollamos en la sección de [bloque]\".",
    ])

    _saltar_pagina(doc)


# =====================================================================
# BLOQUE 2 — DATOS Y EDA
# =====================================================================
def render_bloque_2(doc: Document) -> None:
    add_heading(doc, "3. Bloque 2 · Datos y EDA  (4 min · Fernando)", nivel=1)
    add_mensaje_clave(
        doc,
        "Los rangos óptimos no son un manual copiado — son percentiles "
        "P10/P90 derivados estadísticamente de 25 campañas argentinas reales."
    )

    add_heading(doc, "Slide 2.1 — Las 4 fuentes de datos  (~80 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
LAS 4 FUENTES DE DATOS

| Fuente      | Variable                    |
|-------------|-----------------------------|
| MAGyP       | Producción y rendimientos   |
|             | 11 cultivos × 30 deptos     |
|             | × 25 campañas               |
| Open-Meteo  | Reanálisis ERA5             |
|             | Lluvia · Temp · Heladas     |
| SoilGrids   | Propiedades del suelo       |
|             | pH · MO · Arcilla · Arena   |
| INTA / FAO  | Conocimiento agronómico     |
|             | Ciclos · sensibilidades     |

→ 3.786 registros georreferenciados""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "Empezamos por integrar 4 fuentes oficiales en un dataset maestro único.",
        "MAGyP es la fuente principal — Ministerio de Agricultura, datos "
        "públicos. Tiene 25 campañas históricas de producción por "
        "departamento argentino. Cubre 11 cultivos extensivos.",
        "Open-Meteo es la API de reanálisis climático ERA5 — clima histórico "
        "georreferenciado. Lluvia, temperatura media, días de helada.",
        "SoilGrids es de ISRIC, base global de suelos. pH, materia orgánica, "
        "contenido de arcilla y arena.",
        "La última fuente es conocimiento agronómico estructurado — ciclos de "
        "cultivo, sensibilidades a heladas, recomendaciones de cultivo "
        "predecesor. Lo cargamos a mano en un archivo Prolog desde fuentes "
        "del INTA, FAO y la propia cátedra.",
        "El pipeline en Python consolida todo en 3.786 registros "
        "georreferenciados.",
    ])

    add_heading(
        doc,
        "Slide 2.2 — Cómo derivamos los rangos óptimos  (~90 seg)  ⭐",
        nivel=2,
    )
    add_slide(doc, "Contenido del slide", """\
LOS RANGOS ÓPTIMOS NO ESTÁN HARDCODEADOS

Para cada cultivo, en cada región productiva, filtramos las
campañas EXITOSAS (rendimiento ≥ P25):

     Soja en Pampeana
     campañas exitosas (155 registros)
              |
              v
     pH:    P10 = 6.18    P90 = 6.96
     MO:    P10 = 2.42    P90 = 4.44
     Arc:   P10 = 18.50   P90 = 36.07
              |
              v
     rango_optimo(soja, ph, 6.18, 6.96).
     rango_optimo(soja, mo, 2.42, 4.44).
     rango_optimo(soja, arcilla, 18.5, 36.07).

→ 265 hechos generados desde el dataset
→ Auditables contra los datos""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "Acá quiero detenerme un segundo, porque es uno de los puntos que "
        "considero más sólidos del trabajo.",
        "En un sistema experto típico, los rangos óptimos los pone el experto "
        "humano — son números de manual.",
        "Nosotros tomamos otro camino: derivamos los rangos estadísticamente "
        "desde el dataset.",
        "Para cada cultivo, en cada región productiva, filtramos solamente "
        "las campañas exitosas (las que tuvieron rendimiento por encima del "
        "percentil 25 zonal). Sobre ese subconjunto calculamos los "
        "percentiles 10 y 90 de cada variable agronómica.",
        "Por ejemplo, para soja en Pampeana: 155 campañas exitosas dieron un "
        "rango de pH de 6.18 a 6.96. Esos números están en los datos, no los "
        "inventamos.",
        "El generador automático corre el pipeline y produce 265 hechos "
        "Prolog tipo rango_optimo(soja, ph, 6.18, 6.96).",
        "¿Por qué importa? Porque cualquier evaluador puede tomar el dataset, "
        "correr el notebook, y verificar que esos números son los percentiles "
        "10 y 90 de las campañas exitosas argentinas. Las reglas son "
        "auditables contra los datos.",
    ])

    add_heading(doc, "Slide 2.3 — Cobertura geográfica  (~60 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
COBERTURA GEOGRÁFICA: 3 REGIONES PRODUCTIVAS

  ✅ Pampeana    → BA, SF, CO, ER, La Pampa
  ✅ NOA         → Salta, Tucumán, S.del Estero
  ✅ NEA         → Corrientes, Chaco, Formosa, Misiones

  ❌ Cuyo        → MAGyP no reporta cereales bajo riego
  ❌ Patagonia   → Avena para verdeo no extensiva

  ⚠ Limitación de la fuente, no del sistema.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "Con el alcance del trabajo cubrimos las 3 regiones productivas "
        "argentinas más relevantes: Pampeana, NOA y NEA.",
        "Cuyo y Patagonia quedaron afuera. Es importante decir por qué: "
        "MAGyP no reporta como cultivo extensivo los cereales bajo riego en "
        "Mendoza, ni la avena para verdeo patagónica.",
        "No es que el sistema no las pueda manejar — es que no hay datos en "
        "la fuente oficial para entrenar.",
        "Es una limitación que reconocemos abiertamente, no la disimulamos. "
        "El sistema cubre lo que el MAGyP cubre.",
    ])

    add_frase_ancla(
        doc,
        "Transición al bloque 3 (Fernando → Ezequiel)",
        "Sobre estos datos construimos las dos capas del sistema híbrido. "
        "Voy a dejarle la palabra a Ezequiel para que les cuente la primera, "
        "que es el sistema experto en Prolog."
    )

    _saltar_pagina(doc)


# =====================================================================
# BLOQUE 3 — SISTEMA EXPERTO PROLOG (EZE)
# =====================================================================
def render_bloque_3(doc: Document) -> None:
    add_heading_eze(
        doc,
        "4. Bloque 3 · Sistema experto Prolog  (6 min · EZEQUIEL)  ⭐",
        nivel=1,
    )
    add_mensaje_clave(
        doc,
        "El razonamiento simbólico es auditable, trazable y se puede "
        "inspeccionar manualmente. Cada inferencia tiene un porqué."
    )

    add_heading(doc, "Slide 3.1 — Por qué Prolog  (~70 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
¿POR QUÉ UN SISTEMA EXPERTO EN PROLOG?

Random Forest:    "soja, 3.359 kg/ha, 87% confianza"
                  → Caja negra. Sin trazabilidad.

Sistema experto:  apto(lote_pergamino, soja). → true.
                  → Inferencia auditable, paso a paso.

Cubre el "POR QUÉ" que RF no puede dar.

Tecnología:  SWI-Prolog 9.2.9 + pyswip
             5 archivos de reglas
             265 hechos derivados estadísticamente""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets_eze(doc, [
        "Gracias Fer. Mi parte cubre la capa simbólica del sistema, que es "
        "donde la materia \"Fundamentos de IA\" aparece más explícita.",
        "¿Por qué un sistema experto en Prolog y no solo Random Forest? "
        "Porque Random Forest es una caja negra eficaz pero opaca.",
        "El sistema experto Prolog hace lo opuesto: cada inferencia es "
        "trazable paso a paso.",
        "Tecnología: SWI-Prolog 9.2.9 con pyswip para los bindings con "
        "Python. Base de conocimiento en 5 archivos de reglas, más 265 "
        "hechos generados desde el dataset que mostró Fer.",
    ])

    add_heading(
        doc,
        "Slide 3.2 — Anatomía de una regla  (~100 seg)  ⭐ TÉCNICO CLAVE",
        nivel=2,
    )
    add_slide(doc, "Contenido del slide", """\
EJEMPLO: ¿Es apta la soja para este lote?

apto(Lote, Cultivo) :-
    cultivo_soportado(Cultivo),
    region_lote(Lote, Region),
    cultivo_en_region(Cultivo, Region),
    apto_suelo(Lote, Cultivo),
    apto_clima(Lote, Cultivo).

apto_suelo(Lote, Cultivo) :-
    valor_lote(Lote, ph, V),
    rango_optimo(Cultivo, ph, Min, Max),
    V >= Min, V =< Max,
    %  ... idem MO, arcilla, arena
    .

→ Cada paso es inspeccionable.
→ Si falla, sabemos POR QUÉ falla.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets_eze(doc, [
        "Voy a mostrar el ejemplo más representativo: la regla que decide si "
        "un cultivo es apto para un lote.",
        "La regla apto/2 dice: un cultivo es apto para un lote si se cumplen "
        "5 condiciones simultáneamente — es un AND lógico.",
        "apto_suelo a su vez se descompone: el valor del pH del lote tiene "
        "que estar dentro del rango óptimo del cultivo. Lo mismo para MO, "
        "arcilla, arena.",
        "Lo importante es que cada paso es inspeccionable. Si la consulta "
        "falla, podemos preguntarle a Prolog en qué paso falló.",
        "Es además ejecutable: apto(lote_pergamino, soja) responde true en "
        "1.4 ms. La respuesta no es probabilística — es una verdad lógica "
        "sobre la base de conocimiento.",
    ])

    add_heading(doc, "Slide 3.3 — Los 7 tipos de riesgo  (~90 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
DIAGNÓSTICO SIMBÓLICO DE RIESGOS

  ⚠ Sequía              (lluvia bajo P10 del cultivo)
  ⚠ Helada              (cultivo sensible + DDH > 0)
  ⚠ Acidez              (pH < 5.5)
  ⚠ Alcalinidad         (pH > 8.0)
  ⚠ Exceso hídrico      (lluvia sobre P90 + drenaje malo)
  ⚠ Estrés térmico      (temp media fuera de rango)
  ⚠ Déficit nutricional (MO < umbral del cultivo)

Si hay riesgo crítico:
   recomendado → apto parcial (con manejo correctivo)""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets_eze(doc, [
        "Más allá de la aptitud, el sistema diagnostica 7 tipos de riesgo.",
        "Los 7: Sequía, helada, acidez, alcalinidad, exceso hídrico, estrés "
        "térmico, déficit nutricional. Cada uno es una regla Prolog "
        "independiente.",
        "Por ejemplo, la regla de helada combina dos cosas: que el cultivo "
        "sea sensible a heladas y que el lote tenga días de helada distintos "
        "de cero.",
        "Si hay riesgo crítico, el bridge degrada el cultivo de "
        "\"recomendado\" a \"apto parcial\". No lo descarta, pero le marca al "
        "productor que necesita manejo correctivo.",
        "Este nivel de diagnóstico es típicamente lo que un sistema experto "
        "agronómico haría — y lo implementamos en ~500 líneas de Prolog.",
    ])

    add_heading(doc, "Slide 3.4 — Ejecución en vivo  (~90 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
TODO ESTO ES EJECUTABLE Y AUDITABLE

  ?- apto(lote_pergamino, soja).
     true.                      [1.4 ms]

  ?- riesgo_critico(lote_sequia, soja).
     true.                      [0.8 ms]

  ?- findall(C, recomendar(lote_pergamino, C), L).
     L = [soja, maiz, girasol, sorgo].

→ Tab "Consola Prolog" en la app:
  cualquier evaluador puede ejecutar consultas y ver el sistema
  simbólico responder en vivo.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets_eze(doc, [
        "Para cerrar mi parte: todo esto es ejecutable en vivo.",
        "apto(lote_pergamino, soja) → true. Si pregunto si hay riesgo crítico "
        "de sequía sobre un lote en sequía, responde true también.",
        "findall, meta-predicado de Prolog, enumera todos los cultivos "
        "recomendados para un lote.",
        "Construimos un tab en la app llamado \"Consola Prolog\". Cualquier "
        "evaluador puede entrar a la app, escribir consultas, y ver el "
        "sistema simbólico responder en tiempo real. Las reglas no son algo "
        "que quedó en el código — están expuestas y auditables.",
    ])

    add_frase_ancla(
        doc,
        "Transición al bloque 4 (Ezequiel → Fernando)",
        "Bueno, esa es la capa simbólica. Le devuelvo la palabra a Fer para "
        "que cuente la capa cuantitativa: los modelos Random Forest."
    )

    add_tips(doc, "Tips del bloque 3 (importantes — Eze ensayar)", [
        "El bloque 3 es el más técnico de la presentación. Es donde el "
        "profesor va a prestar más atención. Eze tiene que dominar las 4 "
        "slides — no leer.",
        "La slide 3.2 (regla apto/2) es el momento clave.",
        "Si el profesor interrumpe con una pregunta sobre Prolog, dejá que "
        "Eze la responda — es su sección.",
        "Eze tiene que ensayar leyendo la regla en voz alta. Mejor explicar "
        "el concepto de cada cláusula que leer la sintaxis literal.",
    ])

    _saltar_pagina(doc)


# =====================================================================
# BLOQUE 4 — RANDOM FOREST
# =====================================================================
def render_bloque_4(doc: Document) -> None:
    add_heading(
        doc, "5. Bloque 4 · Modelos Random Forest  (4 min · Fernando)", nivel=1
    )
    add_mensaje_clave(
        doc,
        "Reportamos R² honestos, sin maquillar. El sistema reconoce sus "
        "propios límites en lugar de inflar métricas."
    )

    add_heading(doc, "Slide 4.1 — Por qué Random Forest  (~70 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
LA CAPA CUANTITATIVA: 11 MODELOS RANDOM FOREST

¿Por qué Random Forest?

  · Robusto a outliers (campañas catastróficas)
  · No requiere normalización ni feature scaling
  · Maneja bien la heterogeneidad regional
  · Da intervalos de confianza (varianza entre árboles)

Configuración:
  · 200 árboles · profundidad 15
  · Un modelo INDEPENDIENTE por cultivo
  · Train/test split 80/20

Variable objetivo: rendimiento (kg/ha)
Features: 7 variables agronómicas + región""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "La capa simbólica responde \"qué cultivos son aptos\". La capa "
        "cuantitativa responde la pregunta complementaria: \"cuánto van a "
        "producir\".",
        "Entrenamos 11 modelos Random Forest, uno independiente por cultivo.",
        "Random Forest fue elección consciente. Robusto a outliers (sequía "
        "2022-23). No requiere feature scaling. Maneja heterogeneidad "
        "regional.",
        "La varianza entre árboles del bosque nos da un intervalo de "
        "confianza sin tener que asumir distribuciones.",
        "Configuración: 200 árboles, profundidad máxima 15, split 80/20 "
        "entre train y test. Features: 7 variables agronómicas más región "
        "codificada.",
    ])

    add_heading(
        doc,
        "Slide 4.2 — R² por cultivo  (~120 seg)  ⭐ HONESTIDAD",
        nivel=2,
    )
    add_slide(doc, "Contenido del slide", """\
RESULTADOS — REPORTADOS HONESTAMENTE

  🟢 BUEN AJUSTE (R² ≥ 0.6)
     Algodón     0.818      Maíz        0.638
     Trigo       0.747      Sorgo       0.626
     Centeno     0.741
     Cebada      0.710

  🟡 AJUSTE MODERADO (0.3 ≤ R² < 0.6)
     Soja        0.577      Arroz       0.437
     Avena       0.480      Girasol     0.320

  🔴 DRIVERS FUERA DEL DATASET (R² < 0.3)
     Maní        0.024  ←  el sistema reconoce su límite

→ No inflamos R². Reportamos lo que dieron los datos.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "11 modelos, R² de 0.818 a 0.024. Reportamos así, sin agruparlos para "
        "hacer un promedio que oculte la dispersión.",
        "6 cultivos con buen ajuste — R² > 0.6. Algodón 0.818, trigo 0.747. "
        "Sólido para problema agrícola, donde hay variabilidad estocástica.",
        "4 cultivos en ajuste moderado, 0.3 a 0.6. Sirven como referencia "
        "pero hay que comunicar incertidumbre alta al productor.",
        "Maní en 0.024 — prácticamente sin capacidad predictiva.",
    ])

    add_heading(doc, "Slide 4.3 — Honestidad metodológica  (~50 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
¿POR QUÉ MANÍ TIENE R² = 0.024?

Los drivers reales del rendimiento de maní
NO ESTÁN en nuestro dataset:

  · Variedad genética
  · Fecha exacta de siembra
  · Calcio del suelo
  · Manejo del cultivo

Decisión: NO usar técnicas que inflen el R²
artificialmente. El sistema documenta esta
limitación en lugar de disimularla.

→ Honestidad metodológica.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "Maní tan mal porque los drivers reales no están en nuestro dataset. "
        "Variedad genética. Fecha de siembra. Calcio del suelo. Manejo "
        "específico.",
        "Decisión metodológica: no usar técnicas que inflaran el R² "
        "artificialmente. Hubiera sido maquillaje.",
        "El sistema documenta abiertamente que para maní no puede dar "
        "predicciones confiables.",
        "\"Honestidad metodológica\" como valor del trabajo, no falla.",
    ])

    _saltar_pagina(doc)


# =====================================================================
# BLOQUE 5 — BRIDGE EN CASCADA
# =====================================================================
def render_bloque_5(doc: Document) -> None:
    add_heading(
        doc, "6. Bloque 5 · Bridge en cascada  (5 min · Fernando)  ⭐⭐⭐", nivel=1
    )
    add_mensaje_clave(
        doc,
        "Las dos capas no compiten, cooperan. Prolog filtra, Random Forest "
        "cuantifica, Prolog audita. Esa cascada es nuestro aporte propio."
    )
    add_caption(
        doc,
        "ESTE BLOQUE DEFINE LA NOTA. Es el aporte propio del trabajo. Si hay "
        "que recortar tiempo, recortar en bloque 4 (modelos), no acá."
    )

    add_heading(doc, "Slide 5.1 — La pregunta de diseño  (~60 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
¿CÓMO COMBINAR LAS DOS CAPAS?

Opción 1 — Random Forest puro:
  → Predicción opaca. "Soja, 3.359 kg/ha, 87%."
  → No explica por qué soja, qué hace que sea apta,
     qué riesgos enfrenta.

Opción 2 — Sistema experto puro:
  → No cuantifica el rendimiento esperado.

Opción 3 — Cascada híbrida:  ← LA QUE ELEGIMOS
  → Razonamiento simbólico FILTRA candidatos
  → Random Forest CUANTIFICA solo a los aptos
  → Razonamiento simbólico AUDITA con riesgos

→ Decisión simultáneamente EXPLICABLE y CUANTIFICABLE.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "Tres opciones de arquitectura. Solo Random Forest: opaco. Solo "
        "sistema experto: no cuantifica. Cascada híbrida: la que elegimos.",
        "La cascada es nuestro aporte propio del trabajo.",
    ])

    add_heading(
        doc,
        "Slide 5.2 — La cascada en 3 etapas  (~150 seg)  ⭐⭐ CRÍTICA",
        nivel=2,
    )
    add_slide(doc, "Contenido del slide", """\
LA CASCADA EN 3 ETAPAS

  ETAPA 1 — APTITUD (Prolog)
  ┌────────────────────────────────────────────┐
  │ Para cada cultivo de los 11:               │
  │   ¿es apto este lote?                      │
  │   apto(Lote, Cultivo) → true / false       │
  └────────────────────────────────────────────┘
                      ↓
            cultivos aptos: {soja, maíz, girasol, sorgo}
                      ↓
  ETAPA 2 — PREDICCIÓN (Random Forest)
  ┌────────────────────────────────────────────┐
  │ Solo para los aptos:                       │
  │   modelo_soja.predict(features)            │
  │   modelo_maiz.predict(features)            │
  │   ...                                      │
  │ → predicción + intervalo de confianza      │
  └────────────────────────────────────────────┘
                      ↓
  ETAPA 3 — AUDITORÍA DE RIESGOS (Prolog)
  ┌────────────────────────────────────────────┐
  │ Para cada cultivo apto:                    │
  │   ¿hay riesgo crítico?                     │
  │   Si SÍ → "apto parcial" (manejo correctivo)│
  │   Si NO → "recomendado"                    │
  └────────────────────────────────────────────┘
                      ↓
              REPORTE FINAL al productor""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "La cascada tiene 3 etapas. Voy a recorrerlas una por una.",
        "Etapa 1, aptitud simbólica. Para cada cultivo, Prolog responde true "
        "o false. Si es false, el cultivo queda descartado con motivo "
        "explícito (\"soja no apta porque pH 4.5 fuera de rango óptimo\").",
        "Etapa 2, predicción cuantitativa. Random Forest solo para los "
        "cultivos aptos. Si pasaron 4, corremos 4 modelos. Cada uno devuelve "
        "predicción puntual + intervalo de confianza.",
        "Etapa 3, auditoría simbólica de riesgos. Volvemos a Prolog. Si hay "
        "riesgo crítico, el cultivo se degrada de \"recomendado\" a \"apto "
        "parcial\".",
        "El reporte final que ve el productor sale de esa cascada.",
    ])

    add_heading(doc, "Slide 5.3 — Por qué la cascada importa  (~90 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
¿QUÉ GANA EL SISTEMA CON ESTA ARQUITECTURA?

  1. EFICIENCIA
     No corremos los 11 modelos siempre.
     Solo los que pasaron filtro simbólico.

  2. EXPLICABILIDAD
     Si soja NO es recomendada, podemos decir
     exactamente POR QUÉ:
     "fue descartada en etapa 1 por pH=4.5,
      fuera del rango óptimo [6.18, 6.96]"

  3. AUDITORÍA SOBRE LA PREDICCIÓN
     Si RF dice "soja, 3.500 kg/ha"
     pero hay riesgo crítico de sequía,
     el sistema lo degrada a "apto parcial".
     RF solo nunca podría hacer eso.

→ Las dos paradigmas no compiten. Cooperan.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets(doc, [
        "3 cosas que gana el sistema con la arquitectura híbrida.",
        "1) Eficiencia: no corremos los 11 modelos siempre. Lote pampeano "
        "típico: 4 modelos. Lote ácido fuerte: 0.",
        "2) Explicabilidad: si soja no aparece en el reporte, podemos decir "
        "por qué con regla concreta.",
        "3) Auditoría sobre predicción cuantitativa: si RF dice \"3.500 "
        "kg/ha\" pero hay riesgo crítico de sequía, el sistema degrada el "
        "reporte. RF solo nunca podría hacer eso.",
        "Esto es lo que justifica el sistema híbrido. Las dos paradigmas "
        "centrales de IA, no compitiendo, cooperando.",
    ])

    _saltar_pagina(doc)


# =====================================================================
# BLOQUE 6 — DEMO EN VIVO
# =====================================================================
def render_bloque_6(doc: Document) -> None:
    add_heading(
        doc, "7. Bloque 6 · Demo en vivo  (6 min · Fernando)  ⭐⭐⭐", nivel=1
    )
    add_mensaje_clave(
        doc,
        "Todo lo que dijimos antes funciona en vivo, en producción, y es "
        "reproducible por cualquier evaluador."
    )
    add_caption(
        doc,
        "ALTO RIESGO EN ZOOM. Calentar la app 5 min antes. 2 ventanas abiertas "
        "(producción + local). No improvisar consultas Prolog."
    )

    add_heading(doc, "Plan de 4 escenas", nivel=2)

    add_heading(doc, "Escena 1 — Vista panorámica (60 seg)", nivel=3)
    add_bullets(doc, [
        "Compartir pantalla: https://caece-agrosmart.streamlit.app",
        "Mostrar header CAECE, sidebar con sliders, mapa con 30 "
        "departamentos, los 3 tabs principales.",
    ])

    add_heading(doc, "Escena 2 — Caso típico (90 seg)", nivel=3)
    add_bullets(doc, [
        "Cargar ejemplo \"Pergamino típico\".",
        "Sliders se autocompletan (pH 6.5, MO 3.2%, etc.).",
        "Apretar \"Evaluar lote\".",
        "Mostrar 4 cultivos recomendados (soja, maíz, sorgo, girasol).",
        "Si hay tiempo: detalle de soja con tabla de rangos óptimos + "
        "boxplot de distribución histórica.",
    ])

    add_heading(
        doc,
        "Escena 3 — Consola Prolog en vivo (120 seg)  ⭐⭐⭐ NO SACRIFICAR",
        nivel=3,
    )
    add_bullets(doc, [
        "Cambiar al tab Consola Prolog.",
        "Ejecutar 3 consultas:",
        "  1) apto(lote_pergamino, soja).  → true. (1.4 ms)",
        "  2) apto(lote_acido, soja).      → false.",
        "  3) findall(C, recomendar(lote_pergamino, C), L).",
        "     → L = [soja, maiz, girasol, sorgo]",
        "Cierre: \"Esta consola está disponible para cualquiera. Las reglas "
        "no están escondidas en el código — están auditables en producción.\"",
    ])

    add_heading(doc, "Escena 4 — Caso límite, sequía 2022/23 (90 seg)", nivel=3)
    add_bullets(doc, [
        "Cargar ejemplo \"Sequía 2022/23\".",
        "Lluvia 280 mm anuales (debajo de P10 de todos los cultivos pampeanos).",
        "Click Evaluar lote.",
        "\"Recomendados: ninguno. Aptos parciales: ninguno.\"",
        "\"El sistema reconoce la incertidumbre extrema.\"",
    ])

    add_heading(doc, "Plan B — tabla de fallas posibles", nivel=2)
    add_tabla_simple(
        doc,
        ["Falla", "Plan B"],
        [
            ["App en Streamlit no responde", "Cambiar a localhost"],
            ["Error 500 al cargar ejemplo", "Cargar sliders manualmente"],
            ["Consola Prolog tira error", "Pasar a consultas predefinidas"],
            ["Se cae internet/Zoom", "Pedir 1 min para reconectar"],
            ["No alcanza el tiempo", "Cortar escena 4 (sequía)"],
            ["Sobra tiempo", "Mostrar tab Documentación + diagrama"],
        ],
    )

    add_tips(doc, "Tips críticos de la demo", [
        "Reducí zoom del navegador a 110-125 % antes de compartir pantalla.",
        "No leas la consola palabra por palabra. Tipeá la consulta y contala "
        "mientras tipeás.",
        "Hablá mientras la app calcula. 2 segundos de espera son eternidad si "
        "te quedás callado. Llenalos: \"mientras procesa, esto está corriendo "
        "la cascada de 3 etapas que mencioné\".",
        "Si algo carga lento, no te disculpes. \"Streamlit Cloud está "
        "calentando los caches.\"",
        "No interactúes con cosas que no ensayaste.",
    ])

    add_frase_ancla(
        doc,
        "Transición al bloque 7 (Fernando → Ezequiel)",
        "Esto es lo que construimos y cómo funciona. Le voy a dejar la "
        "palabra a Eze para que cierre con limitaciones y trabajo futuro."
    )

    _saltar_pagina(doc)


# =====================================================================
# BLOQUE 7 — CIERRE (EZE)
# =====================================================================
def render_bloque_7(doc: Document) -> None:
    add_heading_eze(
        doc, "8. Bloque 7 · Cierre  (2 min · EZEQUIEL)", nivel=1
    )
    add_mensaje_clave(
        doc,
        "El sistema reconoce sus propios límites y propone trabajo futuro "
        "concreto. Honestidad metodológica como cierre."
    )

    add_heading(doc, "Slide 7.1 — Limitaciones reconocidas  (~70 seg)", nivel=2)
    add_slide(doc, "Contenido del slide", """\
LIMITACIONES QUE RECONOCEMOS

  · Cobertura geográfica
    Cuyo y Patagonia sin datos. Limitación de MAGyP,
    no del sistema.

  · R² bajo en cultivos minoritarios
    Maní 0.024 — drivers reales fuera del dataset.
    Decidimos NO maquillar las métricas.

  · Granularidad temporal
    El sistema usa promedios de campaña. No incluye
    ventanas críticas (siembra, floración, llenado).

  · Validación agronómica
    "No usar para decisiones reales sin validación
     profesional" — disclaimer explícito en la app.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets_eze(doc, [
        "Gracias Fer. Cierro hablando de las limitaciones reconocidas, que "
        "considero parte del valor del trabajo.",
        "1) Cobertura geográfica. Cuyo y Patagonia sin datos. Si MAGyP "
        "empezara a reportar, el sistema los integraría sin cambios "
        "estructurales.",
        "2) R² bajo en cultivos minoritarios. Maní 0.024. Decidimos no "
        "maquillar las métricas.",
        "3) Granularidad temporal. Promedios de campaña, no ventanas "
        "críticas (siembra, floración, llenado). Está fuera del alcance "
        "actual.",
        "4) El sistema no está validado agronómicamente para uso productivo. "
        "Disclaimer explícito en la app: \"no usar para decisiones reales "
        "sin validación profesional\". Es herramienta académica, no "
        "reemplazo del agrónomo.",
    ])

    add_heading(
        doc, "Slide 7.2 — Trabajo futuro y conclusiones  (~50 seg)", nivel=2
    )
    add_slide(doc, "Contenido del slide", """\
TRABAJO FUTURO

  · Integrar fenología (NDVI, satélites Sentinel-2)
  · Ampliar cobertura: alta montaña, riego, intensivos
  · Incorporar riesgos no abordados: enfermedades, plagas
  · Validar predicciones contra campañas 2025/26 reales

CONCLUSIONES

  · Construimos un sistema híbrido funcional
    con datos argentinos reales

  · Las dos paradigmas de IA cooperan, no compiten

  · El código es reproducible y auditable
    https://github.com/fcaceres-create/caece-agro

  · Honestidad metodológica como criterio rector

  → Gracias. Quedamos abiertos a preguntas.""")
    add_heading(doc, "Bullets de improvisación", nivel=3)
    add_bullets_eze(doc, [
        "Trabajo futuro: integrar fenología satelital (Sentinel-2 da NDVI "
        "cada 5 días); ampliar cobertura a regiones e intensivos; incorporar "
        "riesgos biológicos; validar predicciones contra campañas 2025-26.",
        "Conclusiones: sistema híbrido funcional con datos argentinos reales. "
        "Las dos paradigmas cooperan. Código reproducible y auditable en "
        "GitHub. Honestidad metodológica como criterio rector.",
    ])

    add_frase_ancla(
        doc,
        "Frase literal de cierre",
        "Gracias por su atención, profesor. Quedamos abiertos a preguntas."
    )

    add_tips(doc, "Tips para Eze (cierre)", [
        "El cierre es corto y debe ser firme. 2 min es muy poco; si se "
        "acelera por nervios, queda en 60 seg.",
        "No improvises agradecimientos largos.",
        "Después del \"abiertos a preguntas\", silencio total. No sigas "
        "hablando. Aguantá el silencio.",
        "Si el profesor dice \"muy bien\" sin pregunta, podés cerrar con "
        "\"muchas gracias\" y esperar.",
    ])

    _saltar_pagina(doc)


# =====================================================================
# ANEXO A — PREGUNTAS ANTICIPADAS
# =====================================================================
def render_anexo_a(doc: Document) -> None:
    add_heading(doc, "9. Anexo A · Preguntas anticipadas", nivel=1)
    add_caption(
        doc,
        "7 preguntas que probablemente surjan, con bullets de respuesta. "
        "El responsable está indicado en cada caso."
    )

    preguntas = [
        (
            "P1. \"¿Por qué eligieron Random Forest y no XGBoost o redes "
            "neuronales?\"",
            "Fernando",
            [
                "Elección consciente.",
                "Tres razones: robustez a outliers, no requiere "
                "normalización, intervalo de confianza nativo.",
                "XGBoost daría R² un poco mejor pero a costa de mayor "
                "complejidad de tuning y menor explicabilidad.",
                "Redes neuronales requieren más datos. 3.786 registros "
                "distribuidos en 11 cultivos: el volumen no justifica la "
                "complejidad.",
                "Cuello de botella era la calidad del dataset, no el algoritmo.",
            ],
        ),
        (
            "P2. \"¿Cómo manejan datos faltantes? ¿Qué pasa si SoilGrids no "
            "devuelve pH para un departamento?\"",
            "Fernando",
            [
                "Fallback en cascada para cada variable.",
                "Suelo: si SoilGrids falla, promedio de la región. Si la "
                "región falla, promedio nacional.",
                "Clima: Open-Meteo es estable, prácticamente nunca falla.",
                "100 % de los registros del dataset tienen las 7 features "
                "completas. Marcados con flag de calidad.",
            ],
        ),
        (
            "P3. \"Maní quedó con R² = 0.024. ¿Por qué lo dejaron en el "
            "sistema?\"",
            "Fernando",
            [
                "Decisión consciente: para que el sistema sea honesto sobre "
                "sus límites.",
                "Si lo sacábamos, alguien podría preguntar por maní sin saber "
                "que el sistema no lo cubre.",
                "Al dejarlo, el sistema responde con la predicción + aviso "
                "\"modelo con baja capacidad predictiva\".",
                "Coherente con honestidad metodológica.",
            ],
        ),
        (
            "P4. \"¿Por qué no usaron lógica difusa (fuzzy logic)?\"",
            "Ezequiel",
            [
                "Lo evaluamos. Decisión: priorizar trazabilidad sobre matiz.",
                "Reglas booleanas: cuando una falla, sabemos exactamente cuál.",
                "Fuzzy: \"soja apta al 0.73\" es opaco para el productor.",
                "Donde introducimos matiz: en el bridge en cascada (categoría "
                "\"apto parcial\"), no en las reglas.",
            ],
        ),
        (
            "P5. \"¿Cómo aseguran que las reglas Prolog no tienen "
            "contradicciones?\"",
            "Ezequiel",
            [
                "Tres líneas de defensa.",
                "1) 14 tests automatizados con pytest sobre el bridge. Si una "
                "regla nueva contradice una existente, los tests rompen.",
                "2) Revisión manual de la base de hechos. Los 265 hechos se "
                "generan automáticamente con Python.",
                "3) Sandbox de SWI-Prolog protege la consola libre.",
                "Todos los tests pasan en el último commit.",
            ],
        ),
        (
            "P6. \"Si tuvieran que escalar a producción real, ¿qué "
            "cambiarían?\"",
            "El más fluido en el momento",
            [
                "Tres prioridades.",
                "1) Integrar fenología satelital (Sentinel-2).",
                "2) Base de datos persistente (PostgreSQL).",
                "3) Validación agronómica con expertos antes de uso "
                "productivo.",
                "Infra: Streamlit Cloud no escala más allá de cientos de "
                "usuarios. Producción real → backend separado (FastAPI) + "
                "frontend desacoplado.",
            ],
        ),
        (
            "P7. \"¿Por qué un sistema híbrido y no solo Prolog enriquecido "
            "con datos?\"",
            "Fernando",
            [
                "Pregunta excelente porque toca el corazón del diseño.",
                "Respuesta corta: Prolog no puede aprender patrones no "
                "expresados.",
                "Los rangos óptimos los podemos derivar estadísticamente sí.",
                "Pero la interacción no lineal entre variables (pH × MO × "
                "lluvia simultáneamente) escapa al razonamiento simbólico.",
                "Random Forest captura esas interacciones en los árboles. "
                "Prolog no.",
                "Por eso decimos que las dos paradigmas se complementan: "
                "Prolog razona sobre lo que sabemos explícitamente. Random "
                "Forest aprende lo que está implícito en los datos.",
            ],
        ),
    ]

    for pregunta, responde, bullets in preguntas:
        add_heading(doc, pregunta, nivel=2)
        p = doc.add_paragraph()
        run = p.add_run("Responde: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_CAPTION
        run = p.add_run(responde)
        run.italic = True
        run.font.size = Pt(10)
        if responde == "Ezequiel":
            run.font.color.rgb = COLOR_BLOQUE_EZE
        if responde == "Ezequiel":
            add_bullets_eze(doc, bullets)
        else:
            add_bullets(doc, bullets)

    add_tips(doc, "Tips generales para preguntas", [
        "Si el profesor pregunta algo que no sabés: respondé honestamente. "
        "\"Esa parte no la profundizamos. Una decisión que tomamos por X "
        "razón fue [...]\". Mejor honesto que inventar.",
        "Si pregunta algo ya cubierto: no le digas \"ya lo cubrí\". "
        "Reformulá la respuesta brevemente.",
        "Si la pregunta es para Eze pero responde Fernando (o viceversa): "
        "no es problema. Pero dejen pasar 3 segundos antes de responder.",
        "Si te trabás: pedí clarificación. \"Para asegurarme de entender, "
        "¿la pregunta es sobre [reformulación]?\"",
        "No respondas en automático. Pausá 1-2 segundos, mostrá que estás "
        "pensando, y después respondé.",
    ])

    _saltar_pagina(doc)


# =====================================================================
# ANEXO B — FRASES DE CIERRE
# =====================================================================
def render_anexo_b(doc: Document) -> None:
    add_heading(doc, "10. Anexo B · Frases de cierre (4 escenarios)", nivel=1)
    add_caption(
        doc,
        "Frases literales para los 4 escenarios más comunes después del "
        "cierre formal."
    )

    add_heading(
        doc,
        "Escenario A — Defensa cierra bien, profesor dice \"muy buen trabajo\"",
        nivel=2,
    )
    add_frase_ancla(
        doc,
        "Frase literal",
        "Muchas gracias profesor. Fue un trabajo que disfrutamos mucho hacer."
    )

    add_heading(
        doc,
        "Escenario B — Hubo preguntas duras, hay que cerrar con altura",
        nivel=2,
    )
    add_frase_ancla(
        doc,
        "Frase literal",
        "Gracias profesor por las preguntas, son útiles para pensar la "
        "siguiente versión del trabajo."
    )

    add_heading(doc, "Escenario C — Algo falló en la demo", nivel=2)
    add_frase_ancla(
        doc,
        "Frase literal",
        "Pedimos disculpas por el inconveniente técnico. Toda la app está "
        "disponible en la URL del informe y los logs del repositorio. "
        "Quedamos abiertos a una segunda revisión si fuera útil."
    )

    add_heading(
        doc, "Escenario D — Sobra tiempo, profesor invita a comentar", nivel=2
    )
    add_frase_ancla(
        doc,
        "Frase literal",
        "Si nos permite un comentario final: este trabajo fue para nosotros "
        "una oportunidad concreta de aplicar las dos paradigmas centrales de "
        "la materia a un dominio real. El resultado puede tener limitaciones, "
        "pero el proceso de pensarlo y construirlo nos dejó preguntas que "
        "vamos a llevarnos a otras asignaturas. Gracias."
    )

    _saltar_pagina(doc)


# =====================================================================
# ANEXO C — CHECKLIST
# =====================================================================
def render_anexo_c(doc: Document) -> None:
    add_heading(doc, "11. Anexo C · Checklist del día D", nivel=1)
    add_caption(
        doc,
        "Imprimir o tener abierto en una pestaña aparte el día de la defensa."
    )

    bloques = [
        ("Día anterior (21/05)", [
            "Ensayo completo con Eze cronometrado",
            "App probada en producción y local",
            "Internet probado, backup tethering listo",
            "Laptop y celular cargando al 100 %",
            "Papel con dibujo de la cascada a mano",
            "NO TOCAR EL CÓDIGO",
        ]),
        ("Mañana del 22/05", [
            "App despertada (entrar a la URL)",
            "Lectura única del guion",
            "Comida liviana, café moderado",
            "Zoom configurado 30 min antes",
            "Notificaciones silenciadas en TODO",
            "2 ventanas: producción + local corriendo",
            "Camisa, fondo neutro, luz delante",
        ]),
        ("5 minutos antes", [
            "Voz calentada",
            "Respiración 4-7-8 × 3",
            "Avisé a Eze",
            "Vaso de agua a mano",
        ]),
    ]
    for titulo, items in bloques:
        add_heading(doc, titulo, nivel=2)
        for item in items:
            p = doc.add_paragraph(f"☐  {item}")
            p.paragraph_format.space_after = Pt(2)

    add_heading(doc, "Durante la defensa", nivel=2)
    add_bullets(doc, [
        "Hablar despacio (los nervios aceleran).",
        "Mirar cámara, no slides.",
        "Pausar si me trabo.",
        "3 transiciones críticas con Eze (literales en el cuerpo).",
    ])

    add_heading(doc, "Demo en vivo (escenas y tiempos)", nivel=2)
    add_tabla_simple(
        doc,
        ["Escena", "Tiempo", "Crítico"],
        [
            ["1. Panorámica", "60 seg", ""],
            ["2. Pergamino típico", "90 seg", ""],
            ["3. Consola Prolog", "120 seg", "⭐ NO SACRIFICAR"],
            ["4. Sequía 2022/23", "90 seg", ""],
        ],
    )

    add_heading(doc, "Cierre", nivel=2)
    add_frase_ancla(
        doc,
        "Frase literal de cierre",
        "Gracias por su atención, profesor. Quedamos abiertos a preguntas."
    )

    _saltar_pagina(doc)


# =====================================================================
# ANEXO D — RESUMEN EJECUTIVO
# =====================================================================
def render_anexo_d(doc: Document) -> None:
    add_heading(doc, "12. Anexo D · Resumen ejecutivo", nivel=1)
    add_caption(
        doc,
        "Una sola página con la estructura final + transiciones críticas "
        "para consultar bajo presión."
    )

    add_heading(doc, "Estructura definitiva — 30 min — Zoom con el profesor", nivel=2)
    add_tabla_simple(
        doc,
        ["Tiempo", "Bloque", "Quién", "Slides / Escenas"],
        [
            ["00–03", "Apertura", "Fernando", "3 slides"],
            ["03–07", "Datos y EDA", "Fernando", "3 slides"],
            ["07–13", "Sistema experto Prolog", "EZEQUIEL", "4 slides"],
            ["13–17", "Modelos Random Forest", "Fernando", "3 slides"],
            ["17–22", "Bridge en cascada", "Fernando", "3 slides"],
            ["22–28", "Demo en vivo", "Fernando", "4 escenas"],
            ["28–30", "Cierre", "EZEQUIEL", "2 slides"],
            ["30+", "Q&A", "Ambos", "—"],
        ],
    )
    add_caption(doc, "Total: ~18 slides + portada.")

    add_heading(doc, "Puntos de evaluación principales (⭐)", nivel=2)
    add_bullets(doc, [
        "Slide 3.2 — Anatomía de regla Prolog (Eze).",
        "Slide 4.2 — R² honestos por cultivo (Fernando).",
        "Slide 5.2 — Cascada en 3 etapas (Fernando).",
        "Demo escena 3 — Consola Prolog en vivo (Fernando).",
    ])

    add_heading(doc, "3 transiciones críticas (literales)", nivel=2)
    add_frase_ancla(
        doc,
        "Fernando → Ezequiel · final del bloque 2",
        "Sobre estos datos construimos las dos capas del sistema híbrido. "
        "Voy a dejarle la palabra a Ezequiel para que les cuente la primera, "
        "que es el sistema experto en Prolog."
    )
    add_frase_ancla(
        doc,
        "Ezequiel → Fernando · final del bloque 3",
        "Bueno, esa es la capa simbólica. Le devuelvo la palabra a Fer para "
        "que cuente la capa cuantitativa: los modelos Random Forest."
    )
    add_frase_ancla(
        doc,
        "Fernando → Ezequiel · final del bloque 6",
        "Esto es lo que construimos y cómo funciona. Le voy a dejar la "
        "palabra a Eze para que cierre con limitaciones y trabajo futuro."
    )


# =====================================================================
# Ensamblado
# =====================================================================
def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configurar_estilos(doc)
    configurar_pagina_y_footer(doc)

    render_portada(doc)
    render_introduccion(doc)
    render_bloque_1(doc)
    render_bloque_2(doc)
    render_bloque_3(doc)
    render_bloque_4(doc)
    render_bloque_5(doc)
    render_bloque_6(doc)
    render_bloque_7(doc)
    render_anexo_a(doc)
    render_anexo_b(doc)
    render_anexo_c(doc)
    render_anexo_d(doc)

    doc.save(SALIDA)
    bytes_kb = SALIDA.stat().st_size / 1024
    print(f"Generado: {SALIDA}  ({bytes_kb:.1f} KB)")


if __name__ == "__main__":
    main()
